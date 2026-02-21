import os
import time
from datetime import datetime
from contextlib import asynccontextmanager
from typing import Optional, List, Dict, Any
from collections import defaultdict
from threading import Lock

import structlog
import psycopg2
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, BackgroundTasks, Request, Depends
from fastapi.exceptions import RequestValidationError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, field_validator
from slowapi.extension import Limiter
from slowapi import _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded


# Load environment variables
load_dotenv()

# Internal Imports
try:
    from db import (
        session_maker, 
        init_db, 
        User, 
        BrandDocument, 
        ReviewerLearning,
        check_database_health,
        get_pool_status,
        hash_password,
        verify_password,
        create_access_token,
        get_current_user,
        get_db_session
    )
    from deployer import (
        run_generation_with_learning,
        update_generation_with_human_feedback,
        verify_learning_loop,
        BrandLearningMemory,
        groq_circuit_breaker,
        tavily_circuit_breaker,
        db_circuit_breaker,
        logger
    )
except ImportError:
    from db import (
        session_maker, 
        init_db, 
        User, 
        BrandDocument, 
        ReviewerLearning,
        check_database_health,
        get_pool_status,
        hash_password,
        verify_password,
        create_access_token,
        get_current_user,
        get_db_session
    )
    from deployer import (
        run_generation_with_learning,
        update_generation_with_human_feedback,
        verify_learning_loop,
        BrandLearningMemory,
        groq_circuit_breaker,
        tavily_circuit_breaker,
        db_circuit_breaker,
        logger
    )


# ===== REQUEST LOGGING MIDDLEWARE =====

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan management"""
    # Startup
    logger.info("Starting BrandGuard AI API")
    try:
        init_db()
        logger.info("Database initialized")
        
        # Check database health
        if check_database_health():
            logger.info("Database health check passed")
        else:
            logger.warning("Database health check failed")
    except Exception as e:
        logger.error(f"Startup failed: {e}", exc_info=True)
        raise

    yield  # Application runs

    # Shutdown
    logger.info("Shutting down BrandGuard AI API")
    logger.info(f"Final pool status: {get_pool_status()}")


# ===== FASTAPI APP =====

app = FastAPI(
    title="BrandGuard AI API",
    description="Production-ready marketing content generation with brand learning",
    version="2.0.0",
    lifespan=lifespan
)

# Initialize slowapi Limiter
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure for your domains in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Request logging middleware
@app.middleware("http")
async def log_requests(request: Request, call_next):
    """Log all requests with timing"""
    start_time = time.time()
    
    # Log request
    logger.info(
        "request_started",
        method=request.method,
        path=request.url.path,
        client=request.client.host if request.client else "unknown"
    )
    
    try:
        response = await call_next(request)
        process_time = time.time() - start_time
        
        # Log response
        logger.info(
            "request_completed",
            method=request.method,
            path=request.url.path,
            status_code=response.status_code,
            duration_ms=round(process_time * 1000, 2)
        )
        
        response.headers["X-Process-Time"] = str(process_time)
        return response
        
    except Exception as e:
        process_time = time.time() - start_time
        logger.error(
            "request_failed",
            method=request.method,
            path=request.url.path,
            error=str(e),
            duration_ms=round(process_time * 1000, 2),
            exc_info=True
        )
        raise


# Static files
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")


# ===== PYDANTIC MODELS WITH VALIDATION =====

class GenerateRequest(BaseModel):
    business_id: str
    topic: str
    format_type: str = Field(default="Blog Article", alias="format")
    voice: str = "formal"
    
    @field_validator('business_id')
    @classmethod
    def validate_business_id(cls, v):
        if not v or len(v) < 3:
            raise ValueError("business_id must be at least 3 characters")
        return v
    
    @field_validator('topic')
    @classmethod
    def validate_topic(cls, v):
        if not v or len(v) < 5:
            raise ValueError("topic must be at least 5 characters")
        return v


class FeedbackRequest(BaseModel):
    generation_id: str
    business_id: str
    human_approved: bool
    human_score: float = Field(ge=0, le=10)
    human_feedback: str
    
    @field_validator('human_feedback')
    @classmethod
    def validate_feedback(cls, v):
        if not v or len(v) < 10:
            raise ValueError("feedback must be at least 10 characters")
        return v


class GenerateResponse(BaseModel):
    status: str
    generation_id: str
    content: str
    auto_score: float
    creative_angle: str
    topic: str
    format_type: str
    business_id: str
    timestamp: str
    processing_time_ms: Optional[float] = None


class FeedbackResponse(BaseModel):
    success: bool
    message: str
    generation_id: str
    agent_accuracy: Optional[float] = None


class HealthResponse(BaseModel):
    status: str
    timestamp: str
    database: dict
    circuit_breakers: dict
    connection_pool: dict


class RegisterRequest(BaseModel):
    email: str
    business_id: str
    business_name: Optional[str] = None
    password: str

    @field_validator('password')
    @classmethod
    def validate_password(cls, v):
        if len(v) < 8:
            raise ValueError('Password must be at least 8 characters')
        if len(v.encode('utf-8')) > 72:
            raise ValueError('Password must be no more than 72 bytes (UTF-8). Try a shorter password.')
        return v

    @field_validator('business_id')
    @classmethod
    def validate_business_id(cls, v):
        if not v or len(v) < 3:
            raise ValueError('business_id must be at least 3 characters')
        return v


class LoginRequest(BaseModel):
    email: str
    password: str

    @field_validator('password')
    @classmethod
    def validate_password(cls, v):
        if len(v.encode('utf-8')) > 72:
            raise ValueError('Password must be no more than 72 bytes')
        return v


class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    business_id: str
    business_name: Optional[str] = None


# ===== ROUTES =====

@app.get("/login")
async def serve_login():
    """Serve the login/register page"""
    if os.path.exists("login.html"):
        return FileResponse("login.html")
    return JSONResponse(status_code=404, content={"message": "Login page not found"})


@app.get("/")
async def serve_landing():
    """Serve the landing page (entry point)"""
    if os.path.exists("landing_page.html"):
        return FileResponse("landing_page.html")
    return {
        "name": "BrandGuard AI API",
        "version": "2.0.0",
        "status": "operational",
        "docs": "/docs"
    }


@app.get("/app")
async def serve_index():
    """Serve the main app frontend"""
    if os.path.exists("index.html"):
        return FileResponse("index.html")
    return JSONResponse(status_code=404, content={"message": "App page not found"})


@app.get("/feedback")
async def serve_feedback_page():
    """Serve the feedback page"""
    if os.path.exists("feedback.html"):
        return FileResponse("feedback.html")
    return JSONResponse(status_code=404, content={"message": "Feedback page not found"})


# ===== AUTH ENDPOINTS =====

@app.post("/api/auth/register", response_model=TokenResponse)
async def register(body: RegisterRequest):
    """Register a new user account"""
    with get_db_session() as session:
        # Check email uniqueness
        existing_email = session.query(User).filter_by(email=body.email).first()
        if existing_email:
            raise HTTPException(status_code=400, detail="Email already registered")

        # Check business_id uniqueness
        existing_biz = session.query(User).filter_by(business_id=body.business_id).first()
        if existing_biz:
            raise HTTPException(status_code=400, detail="Business ID already taken")

        user = User(
            email=body.email,
            business_id=body.business_id,
            business_name=body.business_name or body.business_id,
            industry='Marketing',
            password_hash=hash_password(body.password)
        )
        session.add(user)
        session.commit()
        session.refresh(user)

        token = create_access_token({"sub": user.business_id})
        logger.info(f"New user registered: {user.business_id}")
        return TokenResponse(
            access_token=token,
            business_id=user.business_id,
            business_name=user.business_name
        )


@app.post("/api/auth/login", response_model=TokenResponse)
async def login(body: LoginRequest):
    """Login with email + password, returns JWT"""
    with get_db_session() as session:
        user = session.query(User).filter_by(email=body.email).first()
        if not user:
            raise HTTPException(status_code=401, detail="Invalid email or password")
        if not user.password_hash:
            raise HTTPException(status_code=401, detail="Account has no password set. Please register again.")
        if not verify_password(body.password, user.password_hash):
            raise HTTPException(status_code=401, detail="Invalid email or password")
        if not user.is_active:
            raise HTTPException(status_code=403, detail="Account is inactive")

        token = create_access_token({"sub": user.business_id})
        logger.info(f"User logged in: {user.business_id}")
        return TokenResponse(
            access_token=token,
            business_id=user.business_id,
            business_name=user.business_name
        )


@app.get("/api/auth/me")
async def get_me(current_user: User = Depends(get_current_user)):
    """Get current authenticated user info"""
    return {
        "id": current_user.id,
        "email": current_user.email,
        "business_id": current_user.business_id,
        "business_name": current_user.business_name,
        "subscription_tier": current_user.subscription_tier,
        "is_active": current_user.is_active,
        "monthly_generation_limit": current_user.monthly_generation_limit,
        "monthly_generations_used": current_user.monthly_generations_used,
    }


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """
    Comprehensive health check endpoint.
    Checks database, circuit breakers, and connection pool.
    """
    timestamp = datetime.utcnow().isoformat()
    
    try:
        # Check database
        try:
            db_healthy = check_database_health()
            db_status = {
                "healthy": db_healthy,
                "pool": get_pool_status()
            }
        except Exception as e:
            db_status = {
                "healthy": False,
                "error": str(e),
                "pool": {"status": "error"}
            }
        
        # Check circuit breakers (library based)
        try:
            circuit_status = {
                "groq": "open" if groq_circuit_breaker.opened else "closed",
                "tavily": "open" if tavily_circuit_breaker.opened else "closed",
                "database": "open" if db_circuit_breaker.opened else "closed"
            }
        except Exception as e:
            logger.warning(f"Error checking circuit breakers: {e}")
            circuit_status = {"status": "error"}
        
        # Overall status
        overall_healthy = db_status.get("healthy", False) and all(
            state == "closed" for state in circuit_status.values()
            if isinstance(state, str)
        )
        
        return {
            "status": "healthy" if overall_healthy else "degraded",
            "timestamp": timestamp,
            "database": db_status,
            "circuit_breakers": circuit_status,
            "connection_pool": get_pool_status()
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return JSONResponse(
            status_code=200, # Still return 200 but with error info
            content={
                "status": "degraded",
                "timestamp": timestamp,
                "error": str(e),
                "database": {"healthy": False},
                "circuit_breakers": {},
                "connection_pool": {}
            }
        )


@app.post("/api/upload")
@limiter.limit("10/minute")
async def upload_documents(
    request: Request,
    files: List[UploadFile] = File(...),
    business_id: str = Form(...),
    content_type: str = Form(...)
):
    """
    Upload brand documents with validation and error handling.
    
    - **files**: List of files to upload
    - **business_id**: Business identifier
    - **content_type**: Type (blog/social/ad)
    """
    try:
        # Validate content_type
        if content_type not in ['blog', 'social', 'ad']:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid content_type: {content_type}. Must be blog/social/ad"
            )
        
        # Validate file count
        if len(files) > 50:
            raise HTTPException(
                status_code=400,
                detail="Maximum 50 files per upload"
            )

        # Get or create user
        with session_maker() as session:
            user = session.query(User).filter_by(business_id=business_id).first()
            if not user:
                user = User(
                    email=f"{business_id}@brandguard.ai",
                    business_id=business_id,
                    business_name=business_id,
                    industry='Marketing'
                )
                session.add(user)
                session.commit()
                session.refresh(user)

        # Create directory
        folder_mapping = {
            'blog': f'brand_blogs/{business_id}',
            'social': f'brand_socials/{business_id}',
            'ad': f'brand_ads/{business_id}'
        }

        target_folder = folder_mapping[content_type]
        os.makedirs(target_folder, exist_ok=True)

        uploaded_files = []
        total_size = 0

        # Save files
        with session_maker() as session:
            for file in files:
                if file.filename:
                    # Validate file size (10MB limit)
                    content = await file.read()
                    if len(content) > 10 * 1024 * 1024:
                        raise HTTPException(
                            status_code=400,
                            detail=f"File {file.filename} exceeds 10MB limit"
                        )
                    
                    total_size += len(content)
                    
                    filename = file.filename
                    filepath = os.path.join(target_folder, filename)

                    # Save to filesystem (optional, for backward compatibility)
                    with open(filepath, 'wb') as f:
                        f.write(content)

                    file_size = os.path.getsize(filepath)
                    
                    # Extract text content for database storage
                    file_content_text = None
                    try:
                        # Try extraction based on file extension
                        if filename.lower().endswith('.pdf'):
                            # Synchronous PDF extraction
                            try:
                                import pdfplumber
                                with pdfplumber.open(filepath) as pdf:
                                    file_content_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                            except ImportError:
                                try:
                                    import PyPDF2
                                    with open(filepath, 'rb') as pdf_file:
                                        reader = PyPDF2.PdfReader(pdf_file)
                                        file_content_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                                except Exception:
                                    pass
                        
                        elif filename.lower().endswith('.docx'):
                            try:
                                import docx
                                doc = docx.Document(filepath)
                                file_content_text = "\n".join(para.text for para in doc.paragraphs)
                            except ImportError:
                                pass

                        # Fallback to direct decoding if extraction failed or for other types
                        if not file_content_text:
                            try:
                                file_content_text = content.decode('utf-8')
                            except UnicodeDecodeError:
                                file_content_text = content.decode('latin-1')

                    except Exception as e:
                        logger.warning(f"Could not extract/decode {filename}: {e}")
                        file_content_text = None

                    # CRITICAL: Sanitize NUL characters for PostgreSQL
                    if file_content_text:
                        file_content_text = file_content_text.replace('\u0000', '')

                    # Save to database WITH sanitized content
                    doc = BrandDocument(
                        user_id=user.id,
                        business_id=business_id,
                        filename=filename,
                        content_type=content_type,
                        file_path=filepath,
                        file_size_bytes=file_size,
                        file_content=file_content_text,  # Store sanitized content in DB
                        status='uploaded'
                    )
                    session.add(doc)
                    uploaded_files.append(filename)

            session.commit()

        logger.info(f"Uploaded {len(uploaded_files)} files ({total_size/1024:.1f}KB) for {business_id}")

        return {
            "success": True,
            "message": f"Uploaded {len(uploaded_files)} files",
            "files": uploaded_files,
            "business_id": business_id,
            "content_type": content_type,
            "total_size_kb": round(total_size / 1024, 2)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@limiter.limit("5/minute")
@app.post("/api/generate", response_model=GenerateResponse)
async def generate_content(request: Request, body: GenerateRequest):
    start_time = time.time()
    
    try:
        logger.info(f"Generation request: {body.business_id} - {body.topic}")
        
        result_dict, generation_id = run_generation_with_learning(
            business_id=body.business_id,
            topic=body.topic,
            format_type=body.format_type,
            voice=body.voice
        )
        
        # CRITICAL VALIDATION: Check content before returning
        if not result_dict.get("content"):
            logger.error("Empty content in result_dict!")
            raise ValueError("Content generation returned empty content")
        
        content_length = len(result_dict["content"].strip())
        if content_length < 50:
            logger.error(f"Content too short: {content_length} chars")
            raise ValueError(f"Generated content too short: {content_length} chars")
        
        processing_time = round((time.time() - start_time) * 1000, 2)
        
        logger.info(f"Generation OK: {generation_id} ({content_length} chars, {processing_time}ms)")
        
        return GenerateResponse(
            status="success",
            generation_id=generation_id,
            content=result_dict["content"],
            auto_score=result_dict["auto_score"],
            creative_angle=result_dict["creative_angle"],
            topic=body.topic,
            format_type=body.format_type,
            business_id=body.business_id,
            timestamp=datetime.utcnow().isoformat(),
            processing_time_ms=processing_time
        )
        
    except Exception as e:
        logger.error(f"Generation failed: {e}", exc_info=True)
        
        if "Circuit breaker OPEN" in str(e):
            raise HTTPException(
                status_code=503,
                detail="Service temporarily unavailable. Please try again in a moment."
            )
        
        raise HTTPException(
            status_code=500, 
            detail=f"Generation failed: {str(e)}"
        )

@app.post("/api/feedback", response_model=FeedbackResponse)
@limiter.limit("5/minute")
async def submit_feedback(request: FeedbackRequest):
    """
    Submit human feedback with validation and error handling.
    """
    try:
        logger.info(f"Feedback for {request.generation_id}")

        success = update_generation_with_human_feedback(
            generation_id=request.generation_id,
            business_id=request.business_id,
            human_approved=request.human_approved,
            human_score=request.human_score,
            human_feedback=request.human_feedback
        )

        if not success:
            raise HTTPException(
                status_code=404,
                detail=f"Generation {request.generation_id} not found"
            )

        memory = BrandLearningMemory(business_id=request.business_id)
        stats = memory.get_learning_stats("blog")

        return FeedbackResponse(
            success=True,
            message="Feedback saved successfully",
            generation_id=request.generation_id,
            agent_accuracy=stats.get('agent_accuracy', 0.0)
        )

    except HTTPException:
        raise
    except psycopg2.OperationalError as e:
        logger.error(f"Database connection failed: {e}")
        raise HTTPException(
            status_code=503, 
            detail="Database temporarily unavailable"
        )
    except Exception as e:
        logger.error(f"Feedback submission failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Feedback failed: {str(e)[:200]}"
        )


@app.get("/api/learning/{business_id}/stats")
@limiter.limit("20/minute")
async def get_learning_stats(request:Request, business_id: str, content_type: str = "blog"):
    """Get learning statistics for a business"""
    try:
        memory = BrandLearningMemory(business_id=business_id)
        stats = memory.get_learning_stats(content_type)
        
        return {
            "business_id": business_id,
            "content_type": content_type,
            "stats": stats
        }

    except Exception as e:
        logger.error(f"Stats retrieval failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Stats retrieval failed: {str(e)[:200]}"
        )


@app.get("/api/learning/verify/{business_id}")
@limiter.limit("5/minute")
async def verify_learning(request:Request, business_id: str, content_type: str = "blog"):
    """Verify learning system is working"""
    try:
        stats = verify_learning_loop(business_id, content_type)

        diagnostics = []
        if stats.get('total_generations', 0) == 0:
            diagnostics.append({
                "level": "error",
                "message": "No learnings saved"
            })
        elif stats.get('agent_accuracy', 0) < 50 and stats.get('human_feedback_count', 0) > 5:
            diagnostics.append({
                "level": "warning",
                "message": "Agent accuracy low"
            })
        else:
            diagnostics.append({
                "level": "success",
                "message": "Learning system operational"
            })

        return {
            "business_id": business_id,
            "content_type": content_type,
            "stats": stats,
            "diagnostics": diagnostics
        }

    except Exception as e:
        logger.error(f"Verification failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Verification failed: {str(e)[:200]}"
        )


@app.post("/api/refresh-metrics/{business_id}")
@limiter.limit("5/minute")
async def refresh_metrics(request: Request, business_id: str, content_type: str = "blog"):
    """
    Refresh brand metrics for a business.
    Useful after uploading new documents to ensure the generator uses the latest style.
    """
    try:
        from deployer import BrandMetricsAnalyzer
        analyzer = BrandMetricsAnalyzer(business_id, content_type)
        analyzer.refresh_metrics()
        
        return {
            "success": True,
            "message": f"Metrics refreshed for {business_id}/{content_type}"
        }
    except Exception as e:
        logger.error(f"Metrics refresh failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Refresh failed: {str(e)}"
        )


# ===== ERROR HANDLERS =====

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Log specific validation errors for debugging"""
    # Create serializable version of errors
    errors = []
    for error in exc.errors():
        # Strip 'ctx' if it's not serializable or just stringify it
        e = error.copy()
        if 'ctx' in e:
            e['ctx'] = {k: str(v) for k, v in e['ctx'].items()}
        errors.append(e)

    import json
    # Ultra-prominent logging for the user to find
    print("\n" + "="*50)
    print(f"VALIDATION ERROR: {request.url}")
    try:
        print(json.dumps(errors, indent=2))
    except Exception:
        print(errors)
    print("="*50 + "\n")
    
    logger.error(f"Validation error for {request.url}: {errors}")
    return JSONResponse(
        status_code=422,
        content={
            "error": "validation_error",
            "message": "The data provided is invalid.",
            "details": errors
        }
    )


@app.exception_handler(429)
async def rate_limit_handler(request: Request, exc: HTTPException):
    """Custom rate limit error response"""
    return JSONResponse(
        status_code=429,
        content={
            "error": "rate_limit_exceeded",
            "message": "Too many requests. Please try again in 60 seconds.",
            "retry_after": 60
        }
    )


@app.exception_handler(500)
async def internal_error_handler(request: Request, exc: Exception):
    """Custom 500 error response"""
    logger.error(f"Internal error: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "internal_server_error",
            "message": "An internal error occurred. Please try again later.",
            "timestamp": datetime.utcnow().isoformat()
        }
    )


# ===== STARTUP INFO =====

if __name__ == "__main__":
    import uvicorn
    logger.info("Starting BrandGuard AI API Server")

    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
        reload=False,  # Set to True for development
        log_level="info",
        access_log=True
    )
