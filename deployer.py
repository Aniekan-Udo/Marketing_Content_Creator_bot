import os
import re
import json
import statistics
import time
import argparse
import threading
from pathlib import Path
from datetime import datetime
from functools import wraps
from ratelimit import limits, sleep_and_retry
from typing import Dict, Optional, Any, List, Tuple
from collections import defaultdict, Counter
from threading import Lock, RLock
from urllib.parse import urlparse
import uuid

from pydantic import Field

import structlog
import psycopg2
import psycopg2.extras
from dotenv import load_dotenv
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type, RetryError
from circuitbreaker import circuit

# AI and Data Tools
from crewai import Agent, Task, Crew, Process, LLM
from crewai.tools import BaseTool
from tavily import TavilyClient
from llama_index.core import Settings, SimpleDirectoryReader, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SemanticSplitterNodeParser, SimpleNodeParser
from llama_index.vector_stores.postgres import PGVectorStore
from llama_index.embeddings.fastembed import FastEmbedEmbedding

# Local Imports
from db import (
    ReviewerLearning,
    get_db_session,
    session_maker,
    User,
    check_database_health,
    get_or_create_user as db_get_or_create_user
)

# Load environment variables
load_dotenv()

# ===== LOGGING SETUP =====
structlog.configure(
    processors=[
        structlog.processors.add_log_level,
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.StackInfoRenderer(),
        structlog.processors.format_exc_info,
        structlog.processors.JSONRenderer() if "JSON" in os.getenv("LOG_FORMAT", "console") else structlog.dev.ConsoleRenderer()
    ],
    logger_factory=structlog.stdlib.LoggerFactory(),
    wrapper_class=structlog.stdlib.BoundLogger,
    cache_logger_on_first_use=True,
)

logger = structlog.get_logger("Marketing_Agent")

# ===== SERVICE CONFIGURATION =====
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
POSTGRES_URI = os.getenv("POSTGRES_URI")
POSTGRES_ASYNC_URI = os.getenv("POSTGRES_ASYNC_URI")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

# Validate critical environment variables
if not POSTGRES_URI or not POSTGRES_ASYNC_URI or not GROQ_API_KEY:
    error_msg = "Critical environment variables (POSTGRES_URI, GROQ_API_KEY) missing"
    logger.error(f"Error: {error_msg}")
    raise ValueError(error_msg)

if not TAVILY_API_KEY:
    logger.warning("TAVILY_API_KEY not found - web search disabled")

logger.info("Environment variables validated")

# ===== CUSTOM DECORATORS =====

# Rate limiters (30 calls/min for Groq, 20 for Tavily)
groq_rate_limit = limits(calls=30, period=60)
tavily_rate_limit = limits(calls=20, period=60)

# Circuit breakers
groq_circuit_breaker = circuit(failure_threshold=5, recovery_timeout=60.0)
tavily_circuit_breaker = circuit(failure_threshold=3, recovery_timeout=30.0)
db_circuit_breaker = circuit(failure_threshold=5, recovery_timeout=30.0)

# ===== LOCKS FOR THREAD SAFETY =====
_init_locks: Dict[str, RLock] = {}
_locks_lock = RLock()

def get_init_lock(business_id: str) -> RLock:
    """Get or create a reentrant lock for business_id"""
    with _locks_lock:
        if business_id not in _init_locks:
            _init_locks[business_id] = RLock()
        return _init_locks[business_id]


# ===== LLM CONFIGURATION WITH RETRY AND RATE LIMITING =====

@sleep_and_retry
@groq_rate_limit
@groq_circuit_breaker
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=10),
    reraise=True
)
def get_llm(temperature=0.7):
    """
    Get configured LLM instance with retry logic and rate limiting.
    
    Args:
        temperature: LLM temperature (0.0 to 1.0)
        
    Returns:
        LLM: Configured LLM instance
    """
    if not GROQ_API_KEY:
        logger.error("GROQ_API_KEY is None or empty!")
        raise ValueError("GROQ_API_KEY not found")
    
    try:
        llm = LLM(
            model="groq/llama-3.3-70b-versatile",
            api_key=GROQ_API_KEY,
            temperature=temperature
        )
        logger.info(f"LLM configured: llama-3.3-70b-versatile (temp={temperature})")
        return llm
        
    except Exception as e:
        logger.error(f"Failed to create LLM: {e}", exc_info=True)
        raise# ===== BRAND METRICS ANALYZER =====

class BrandMetricsAnalyzer:
    """
    Analyzes brand documents to extract concrete, measurable metrics.
    Provides ground truth for automatic scoring.
    
    Thread-safe with caching.
    """
    
    def __init__(self, business_id: str, content_type: str):
        self.business_id = business_id
        self.content_type = content_type
        self.metrics = None
        self._lock = Lock()
        self._load_metrics()
    
    def _load_brand_docs(self) -> List[str]:
        """Load raw brand documents from DATABASE first, filesystem as fallback"""
        docs = []
        
        # PRIORITY 1: Load from database
        try:
            logger.info(f"Loading brand docs from database for {self.business_id}/{self.content_type}...")
            with get_db_session() as session:
                from db import BrandDocument
                db_docs = session.query(BrandDocument).filter_by(
                    business_id=self.business_id,
                    content_type=self.content_type
                ).all()
                
                logger.info(f"Found {len(db_docs)} documents in database")
                
                for db_doc in db_docs:
                    if db_doc.file_content:
                        docs.append(db_doc.file_content)
                        logger.info(f"Loaded from DB: {db_doc.filename} ({len(db_doc.file_content)} chars)")
                    else:
                        logger.warning(f"WARNING: No content in DB for: {db_doc.filename}")
        
        except Exception as db_error:
            logger.warning(f"Database document loading failed: {db_error}")
        
        # FALLBACK: Load from filesystem if database is empty (for local dev)
        if not docs:
            logger.info("No documents in database, checking filesystem...")
            data_path = os.path.abspath(f"brand_{self.content_type}s/{self.business_id}")
            
            if not os.path.exists(data_path):
                logger.warning(f"Brand docs path not found: {data_path}")
                return []
            
            try:
                for filename in os.listdir(data_path):
                    filepath = os.path.join(data_path, filename)
                    if not os.path.isfile(filepath):
                        continue
                    
                    try:
                        content = None
                        
                        # AUTOMATED PDF extraction with multiple fallbacks
                        if filename.lower().endswith('.pdf'):
                            # Try pdfplumber first (best quality)
                            try:
                                import pdfplumber
                                text_parts = []
                                with pdfplumber.open(filepath) as pdf:
                                    for page in pdf.pages:
                                        page_text = page.extract_text()
                                        if page_text:
                                            text_parts.append(page_text)
                                content = "\n".join(text_parts)
                                logger.info(f"Extracted PDF with pdfplumber: {filename}")
                            except ImportError:
                                logger.info("pdfplumber not available, trying PyPDF2...")
                            except Exception as e:
                                logger.warning(f"pdfplumber failed: {e}, trying PyPDF2...")
                            
                            # Fallback to PyPDF2
                            if not content:
                                try:
                                    import PyPDF2
                                    text_parts = []
                                    with open(filepath, 'rb') as file:
                                        reader = PyPDF2.PdfReader(file)
                                        for page in reader.pages:
                                            page_text = page.extract_text()
                                            if page_text:
                                                text_parts.append(page_text)
                                    content = "\n".join(text_parts)
                                    logger.info(f"Extracted PDF with PyPDF2: {filename}")
                                except ImportError:
                                    logger.info("PyPDF2 not available, trying pdfminer...")
                                except Exception as e:
                                    logger.warning(f"PyPDF2 failed: {e}, trying pdfminer...")
                            
                            # Final fallback to pdfminer
                            if not content:
                                try:
                                    from pdfminer.high_level import extract_text
                                    content = extract_text(filepath)
                                    logger.info(f"Extracted PDF with pdfminer: {filename}")
                                except ImportError:
                                    logger.warning("No PDF libraries available (pdfplumber, PyPDF2, pdfminer)")
                                except Exception as e:
                                    logger.error(f"All PDF extraction methods failed for {filename}: {e}")
                        
                        # DOCX files
                        elif filename.lower().endswith('.docx'):
                            try:
                                import docx
                                doc = docx.Document(filepath)
                                content = '\n'.join([para.text for para in doc.paragraphs])
                                logger.info(f"Extracted DOCX: {filename}")
                            except ImportError:
                                logger.warning(f"python-docx not installed, skipping {filename}")
                            except Exception as e:
                                logger.error(f"Failed to read DOCX {filename}: {e}")
                        
                        # Text files
                        else:
                            try:
                                with open(filepath, 'r', encoding='utf-8') as f:
                                    content = f.read()
                            except UnicodeDecodeError:
                                logger.warning(f"Could not decode {filename}, trying latin-1")
                                try:
                                    with open(filepath, 'r', encoding='latin-1') as f:
                                        content = f.read()
                                except Exception as e:
                                    logger.error(f"Failed to read {filename}: {e}")
                        
                        # Add if valid
                        if content and content.strip():
                            docs.append(content)
                            logger.info(f"  Loaded from filesystem: {len(content)} chars")
                        
                    except Exception as e:
                        logger.warning(f"Could not read {filename}: {e}")
                        
            except Exception as e:
                logger.error(f"Error listing directory {data_path}: {e}")
        
        logger.info(f"Total documents loaded: {len(docs)}")
        return docs
    
    def _is_pdf_metadata(self, phrase: str) -> bool:
        """
        Detect and filter PDF metadata artifacts.
        ENHANCED: More intelligent detection
        """
        if not phrase or len(phrase) < 2:
            return True
        
        # Check for numbers (common in PDF metadata)
        if any(char.isdigit() for char in phrase):
            return True
        
        # Check for PDF-specific keywords
        pdf_keywords = ['obj', 'type', 'struct', 'elem', 'endobj', 'xref', 
                       'trailer', 'startxref', 'stream', 'endstream', 'font',
                       'encoding', 'basefont', 'subtype']
        words = phrase.lower().split()
        if any(keyword in words for keyword in pdf_keywords):
            return True
        
        # Check if all words are very short (typical of metadata)
        if words and all(len(w) <= 3 for w in words):
            return True
        
        # Check for single-letter patterns
        if re.match(r'^[a-z]\s+[a-z]\s+[a-z]', phrase):
            return True
        
        # ENHANCED: Check if mostly punctuation or special characters
        alpha_ratio = sum(c.isalpha() for c in phrase) / len(phrase) if phrase else 0
        if alpha_ratio < 0.5:
            return True
        
        return False
    
    def _clean_text_for_analysis(self, text: str) -> str:
        """
        Clean text and remove PDF artifacts before analysis.
        ENHANCED: Adaptive filtering based on document statistics
        """
        if not text:
            return ""
        
        # Remove PDF metadata patterns
        text = re.sub(r'\d+\s+\d+\s+obj', '', text)
        text = re.sub(r'<<.*?>>', '', text)
        text = re.sub(r'/\w+\s+', '', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # ENHANCED: Adaptive line filtering
        lines = text.split('\n')
        
        # Calculate document statistics for adaptive thresholds
        line_lengths = [len(line.strip()) for line in lines if line.strip()]
        if line_lengths:
            try:
                median_length = statistics.median(line_lengths)
                # Adaptive threshold: keep lines longer than 25% of median
                min_length = max(20, int(median_length * 0.25))
            except:
                min_length = 20
        else:
            min_length = 20
        
        cleaned_lines = []
        for line in lines:
            line_stripped = line.strip()
            
            # Skip empty or very short lines (adaptive)
            if len(line_stripped) < min_length:
                continue
            
            # Skip lines that are mostly numbers
            if re.match(r'^[\d\s]+$', line_stripped):
                continue
            
            # ENHANCED: Skip lines with very high digit ratio
            digit_ratio = sum(c.isdigit() for c in line_stripped) / len(line_stripped) if line_stripped else 0
            if digit_ratio > 0.5:
                continue
            
            # ENHANCED: Skip lines that look like PDF commands
            if re.match(r'^[\d\s/]+obj|^<<|^>>|^endobj', line_stripped):
                continue
            
            # Keep line if it has reasonable text content
            words = line_stripped.split()
            if words:
                # Must have at least some real words (length > 3)
                real_words = [w for w in words if len(w) > 3 and w.isalpha()]
                if len(real_words) >= 2 or len(real_words) / len(words) > 0.3:
                    cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """
        Extract metrics from a single text with improved filtering.
        NO SIGNATURE CHANGES
        """
        if not text or not text.strip():
            return self._get_empty_metrics()
        
        try:
            # Clean text first (now uses enhanced adaptive cleaning)
            text = self._clean_text_for_analysis(text)
            
            if not text.strip():
                return self._get_empty_metrics()
            
            # Sentence splitting
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
            
            if not sentences:
                return self._get_empty_metrics()
            
            # Word counts for adaptive thresholding
            sentence_lengths = [len(s.split()) for s in sentences]
            median_sentence_len = statistics.median(sentence_lengths) if sentence_lengths else 15
            
            # Use median to filter outliers instead of hardcoded 10
            min_sentence_words = max(5, int(median_sentence_len * 0.3))
            sentences = [s for s in sentences if len(s.split()) >= min_sentence_words]
            
            avg_sentence_length = statistics.mean(sentence_lengths) if sentence_lengths else 0
            
            # Improved paragraph detection
            paragraphs = re.split(r'\n\s*\n', text)
            paragraphs = [p.strip() for p in paragraphs if p.strip()]
            
            if len(paragraphs) < 3:
                # Try single newline split
                paragraphs = re.split(r'\n+', text)
                # Filter out metadata lines (shorter than median/2)
                paragraphs = [p.strip() for p in paragraphs if p.strip() and len(p.split()) >= min_sentence_words]
            
            # CRITICAL: If still 1 big block, use virtual paragraphs to avoid skewed metrics (e.g. 76.6 sent/para)
            if len(paragraphs) <= 2 and len(sentences) > 10:
                logger.info(f"Using virtual paragraphs for density calculation ({len(sentences)} sentences)")
                chunk_size = 4  # Default reasonable paragraph size
                virtual_paras = []
                for i in range(0, len(sentences), chunk_size):
                    virtual_paras.append(" ".join(sentences[i:i+chunk_size]))
                paragraphs = virtual_paras
            
            # Sentences per paragraph with adaptive logic
            sentences_per_para = []
            for para in paragraphs:
                para_sentences = re.split(r'[.!?]+', para)
                para_sentences = [s.strip() for s in para_sentences if s.strip() and len(s.split()) >= min_sentence_words]
                if para_sentences:
                    sentences_per_para.append(len(para_sentences))
            
            if not sentences_per_para:
                avg_sentences_per_para = 3.0
            else:
                # Use median for robustness against multi-page text blocks
                avg_sentences_per_para = statistics.median(sentences_per_para)
                
                # Dynamic warning threshold (3x the expected average of 3-5)
                if avg_sentences_per_para > 25:
                    logger.info(f"Analysis: Document has high density ({avg_sentences_per_para:.1f} sent/para)")
            
            # Detect formatting
            has_bullets = bool(re.search(r'^\s*[-*•]', text, re.MULTILINE))
            has_numbered_lists = bool(re.search(r'^\s*\d+\.', text, re.MULTILINE))
            
            # Improved phrase extraction
            words = text.lower().split()
            phrases = []
            
            for i in range(len(words) - 1):
                for length in [2, 3, 4, 5]:
                    if i + length <= len(words):
                        phrase = ' '.join(words[i:i+length])
                        phrase = re.sub(r'[^\w\s]', '', phrase).strip()
                        
                        if not phrase or len(phrase) < 10:
                            continue
                        
                        if self._is_pdf_metadata(phrase):
                            continue
                        
                        # ENHANCED: Better stopword filtering
                        stopwords = {'the', 'and', 'for', 'with', 'this', 'that', 'from', 'have', 
                                   'been', 'are', 'was', 'were', 'will', 'would', 'could', 'should'}
                        phrase_words = phrase.split()
                        if all(word in stopwords for word in phrase_words):
                            continue
                        
                        # ENHANCED: Require at least one substantive word
                        if not any(len(word) > 4 for word in phrase_words):
                            continue
                        
                        if len(phrase_words) >= 2:
                            phrases.append(phrase)
            
            # Count phrase frequency
            phrase_counts = Counter(phrases)
            
            # Filter: must appear at least 2 times AND not be metadata
            common_phrases = [
                phrase for phrase, count in phrase_counts.most_common(50) 
                if count >= 2 and not self._is_pdf_metadata(phrase)
            ]
            
            # Additional validation
            validated_phrases = []
            for phrase in common_phrases:
                # Must contain at least one word longer than 4 characters
                if any(len(word) > 4 for word in phrase.split()):
                    validated_phrases.append(phrase)
            
            return {
                'sentence_count': len(sentences),
                'avg_sentence_length': avg_sentence_length,
                'paragraph_count': len(paragraphs),
                'avg_sentences_per_para': avg_sentences_per_para,
                'has_bullets': has_bullets,
                'has_numbered_lists': has_numbered_lists,
                'common_phrases': validated_phrases[:20],
                'total_words': len(text.split())
            }
            
        except Exception as e:
            logger.error(f"Error analyzing text: {e}")
            return self._get_empty_metrics()
    
    def _get_empty_metrics(self) -> Dict[str, Any]:
        """Return empty metrics structure - UNCHANGED"""
        return {
            'sentence_count': 0,
            'avg_sentence_length': 0,
            'paragraph_count': 0,
            'avg_sentences_per_para': 0,
            'has_bullets': False,
            'has_numbered_lists': False,
            'common_phrases': [],
            'total_words': 0
        }
    
    def _load_metrics(self):
        """Load and cache brand metrics with thread safety - UNCHANGED"""
        with self._lock:
            docs = self._load_brand_docs()
            
            if not docs:
                logger.warning(f"No brand docs found for {self.business_id}/{self.content_type}")
                self.metrics = self._get_fallback_metrics()
                return
            
            all_metrics = [self._analyze_text(doc) for doc in docs]
            all_metrics = [m for m in all_metrics if m['sentence_count'] > 0]
            
            if not all_metrics:
                logger.warning(f"No valid metrics extracted for {self.business_id}/{self.content_type}")
                self.metrics = self._get_fallback_metrics()
                return
            
            self.metrics = {
                'target_sentence_length': sum(m['avg_sentence_length'] for m in all_metrics) / len(all_metrics),
                'target_sentences_per_para': sum(m['avg_sentences_per_para'] for m in all_metrics) / len(all_metrics),
                'uses_bullets': any(m['has_bullets'] for m in all_metrics),
                'uses_numbered_lists': any(m['has_numbered_lists'] for m in all_metrics),
                'signature_phrases': self._consolidate_phrases([m['common_phrases'] for m in all_metrics]),
                'sample_count': len(docs)
            }

            logger.info(f"   BRAND METRICS LEARNED for {self.business_id}/{self.content_type}:")
            logger.info(f"   Analyzed {self.metrics['sample_count']} documents")
            logger.info(f"   Target sentence length: {self.metrics['target_sentence_length']:.1f} words")
            logger.info(f"   Target sentences/para: {self.metrics['target_sentences_per_para']:.1f}")
            logger.info(f"   Uses bullets: {self.metrics['uses_bullets']}")
            logger.info(f"   Signature phrases found: {len(self.metrics['signature_phrases'])}")
            
            if self.metrics['signature_phrases']:
                logger.info(f"   Sample phrases: {self.metrics['signature_phrases'][:5]}")
    
    def refresh_metrics(self):
        """Refresh metrics when new documents are uploaded"""
        logger.info(f"Refreshing metrics for {self.business_id}/{self.content_type}...")
        self._load_metrics()
        logger.info("Metrics refreshed successfully")
    
    def _consolidate_phrases(self, phrase_lists: List[List[str]]) -> List[str]:
        """
        Find phrases common across multiple documents with better filtering - UNCHANGED
        """
        all_phrases = [phrase for sublist in phrase_lists for phrase in sublist]
        phrase_counts = Counter(all_phrases)
        
        consolidated = [
            phrase for phrase, count in phrase_counts.most_common(30) 
            if count >= 2 and not self._is_pdf_metadata(phrase)
        ]
        
        return consolidated[:25]
    
    def _get_fallback_metrics(self) -> Dict[str, Any]:
        """Fallback metrics if no docs available - UNCHANGED"""
        return {
            'target_sentence_length': 15.0,
            'target_sentences_per_para': 3.0,
            'uses_bullets': False,
            'uses_numbered_lists': False,
            'signature_phrases': [],
            'sample_count': 0
        }
    
    def score_content(self, content: str) -> Dict[str, Any]:
        """
        Score content against brand metrics.
        Thread-safe and handles errors gracefully - UNCHANGED
        """
        if not content or not content.strip():
            logger.warning("Empty content provided for scoring")
            return self._get_zero_scores()
        
        try:
            content_metrics = self._analyze_text(content)
            
            target_len = self.metrics['target_sentence_length']
            actual_len = content_metrics['avg_sentence_length']
            len_deviation = abs(actual_len - target_len) / target_len if target_len > 0 else 1.0
            sentence_length_score = max(0, 10 - (len_deviation * 12))
            
            target_para = self.metrics['target_sentences_per_para']
            actual_para = content_metrics['avg_sentences_per_para']
            para_deviation = abs(actual_para - target_para) / target_para if target_para > 0 else 1.0
            para_structure_score = max(0, 10 - (para_deviation * 10))
            
            format_aligned = True
            if self.metrics['uses_bullets'] != content_metrics['has_bullets']:
                format_aligned = False
            if self.metrics['uses_numbered_lists'] != content_metrics['has_numbered_lists']:
                format_aligned = False
            format_score = 10 if format_aligned else 5
            
            content_lower = content.lower()
            phrases_found = [p for p in self.metrics['signature_phrases'] if p in content_lower]
            phrase_usage_ratio = len(phrases_found) / max(len(self.metrics['signature_phrases']), 1)
            phrase_score = min(10, phrase_usage_ratio * 20)
            
            structure_score = (sentence_length_score + para_structure_score + format_score) / 3
            
            return {
                'structure_score': round(structure_score, 1),
                'phrase_score': round(phrase_score, 1),
                'sentence_length_score': round(sentence_length_score, 1),
                'para_structure_score': round(para_structure_score, 1),
                'format_score': format_score,
                'details': {
                    'target_sentence_length': round(target_len, 1),
                    'actual_sentence_length': round(actual_len, 1),
                    'target_sentences_per_para': round(target_para, 1),
                    'actual_sentences_per_para': round(actual_para, 1),
                    'signature_phrases_available': len(self.metrics['signature_phrases']),
                    'signature_phrases_used': len(phrases_found),
                    'phrases_found': phrases_found[:10]
                }
            }
            
        except Exception as e:
            logger.error(f"Error scoring content: {e}", exc_info=True)
            return self._get_zero_scores()
    
    def _get_zero_scores(self) -> Dict[str, Any]:
        """Return zero scores structure for error cases - UNCHANGED"""
        return {
            'structure_score': 0.0,
            'phrase_score': 0.0,
            'sentence_length_score': 0.0,
            'para_structure_score': 0.0,
            'format_score': 0,
            'details': {
                'target_sentence_length': 0.0,
                'actual_sentence_length': 0.0,
                'target_sentences_per_para': 0.0,
                'actual_sentences_per_para': 0.0,
                'signature_phrases_available': 0,
                'signature_phrases_used': 0,
                'phrases_found': []
            }
        }

    def get_top_examples(self, limit: int = 3) -> str:
        """
        Fetch top 3 raw content examples from the database to use as Few-Shot Training.
        Returns a formatted string of examples.
        """
        try:
            with get_db_session() as session:
                from db import BrandDocument
                # Naive selection: Get longest documents as they likely contain more 'voice'
                # Filter for non-empty content
                docs = session.query(BrandDocument.file_content).filter(
                    BrandDocument.business_id == self.business_id,
                    BrandDocument.content_type == self.content_type,
                    BrandDocument.file_content.isnot(None)
                ).limit(limit).all()
                
                if not docs:
                    return ""

                examples_str = "\n\n".join([f"--- EXAMPLE {i+1} ---\n{doc.file_content[:1500]}..." for i, doc in enumerate(docs) if doc.file_content])
                return examples_str
        except Exception as e:
            logger.warning(f"Failed to fetch top examples: {e}")
            return ""# ===== BRAND LEARNING MEMORY (PRODUCTION-READY) =====

class BrandLearningMemory:
    """
    Persistent memory system using raw SQL on reviewer_learning table.
    
    Features:
    - Connection pooling
    - Retry logic for transient failures
    - Proper transaction management
    - SQL injection protection via parameterized queries
    """
    
    def __init__(self, business_id: str):
        self.business_id = business_id
        self.table_name = "reviewer_learning"
        self._connection_pool = None
    
    @db_circuit_breaker
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((psycopg2.OperationalError, psycopg2.InterfaceError))
    )
    def _get_connection(self):
        """Get database connection with retry logic"""
        try:
            parsed = urlparse(POSTGRES_URI)
            
            # Build connection parameters (Neon-compatible)
            conn_params = {
                'host': parsed.hostname,
                'port': parsed.port or 5432,
                'database': parsed.path.lstrip('/'),
                'user': parsed.username,
                'password': parsed.password,
                'connect_timeout': 10,
                'sslmode': 'require'  # Required for Neon
            }
            
            # Neon pooler doesn't support startup parameters
            if 'pooler' not in parsed.hostname:
                conn_params['options'] = '-c statement_timeout=30000'
            
            conn = psycopg2.connect(**conn_params)
            conn.autocommit = False  # Explicit transaction control
            
            # Set statement timeout after connection for Neon
            if 'pooler' in parsed.hostname:
                with conn.cursor() as cur:
                    cur.execute("SET statement_timeout = '30s'")
            
            return conn
        except Exception as e:
            logger.error(f"Database connection failed: {e}")
            raise
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((psycopg2.OperationalError, psycopg2.DatabaseError))
    )
    def save_learning(self, 
                     generation_id: str,
                     content_type: str,
                     creative_angle: str,
                     generated_content: str,
                     auto_score: float,
                     human_approved: Optional[bool] = None,
                     human_score: Optional[float] = None,
                     human_feedback: str = "",
                     topic: str = "",
                     user_id: int = None,
                     format_type: str = ""):
        """
        Save learning pattern using raw SQL with proper transaction management.
        """
        conn = None
        cursor = None
        
        try:
            conn = self._get_connection()
            cursor = conn.cursor()
            
            # Check if exists
            cursor.execute(f"""
                SELECT id FROM {self.table_name}
                WHERE generation_id = %s
            """, (generation_id,))
            
            existing = cursor.fetchone()
            
            if existing:
                # Update existing record with human feedback
                cursor.execute(f"""
                    UPDATE {self.table_name}
                    SET has_human_feedback = TRUE,
                        human_approved = %s,
                        human_score = %s,
                        human_feedback = %s,
                        agent_correct = (agent_auto_approved = %s)
                    WHERE generation_id = %s
                """, (human_approved, human_score, human_feedback, 
                      human_approved, generation_id))
                logger.info(f"Updated learning: {generation_id}")
            else:
                # Insert new record
                agent_auto_approved = (auto_score >= 8.0)
                has_human_feedback = (human_approved is not None)
                agent_correct = (human_approved == agent_auto_approved) if has_human_feedback else None
                
                features_used = json.dumps({
                    "creative_angle": creative_angle,
                    "auto_score": auto_score
                })
                
                cursor.execute(f"""
                    INSERT INTO {self.table_name}
                    (generation_id, business_id, user_id, topic, content_type, format_type,
                     generated_content, creative_angle, agent_auto_score, 
                     agent_auto_approved, has_human_feedback, human_approved,
                     human_score, human_feedback, agent_correct, features_used)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (generation_id, self.business_id, user_id, topic, content_type, format_type,
                      generated_content, creative_angle, auto_score,
                      agent_auto_approved, has_human_feedback, human_approved,
                      human_score, human_feedback, agent_correct, features_used))
                logger.info(f"Saved new learning: {generation_id}")
            
            conn.commit()
            return True
            
        except Exception as e:
            logger.error(f"Failed to save learning: {e}", exc_info=True)
            if conn:
                conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((psycopg2.OperationalError, psycopg2.DatabaseError))
    )
    def update_with_human_feedback(self, generation_id: str, 
                               human_approved: bool, 
                               human_score: float,
                               human_feedback: str):
        """Update existing learning record with human feedback"""
        conn = None
        cursor = None
        
        try:
            conn = self._get_connection()
            cursor = conn.cursor()

            cursor.execute(f"""
                UPDATE {self.table_name} 
                SET has_human_feedback = TRUE,
                    human_approved = %s,
                    human_score = %s,
                    human_feedback = %s,
                    agent_correct = (agent_auto_approved = %s)
                WHERE generation_id = %s AND business_id = %s
            """, (human_approved, human_score, human_feedback, 
                human_approved, generation_id, self.business_id))
            
            if cursor.rowcount == 0:
                logger.warning(f"No learning found: {generation_id}/{self.business_id}")
                return False
            
            conn.commit()
            logger.info(f"Feedback saved for {generation_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to update feedback: {e}", exc_info=True)
            if conn:
                conn.rollback()
            return False
            
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((psycopg2.OperationalError, psycopg2.DatabaseError))
    )
    def get_approved_patterns(self, content_type: str, limit: int = 10) -> List[Dict[str, Any]]:
        """Retrieve approved patterns using raw SQL"""
        conn = None
        cursor = None
        
        try:
            conn = db_circuit_breaker.call(self._get_connection)
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cursor.execute(f"""
                SELECT 
                    generation_id,
                    creative_angle,
                    generated_content,
                    agent_auto_score as auto_score,
                    human_approved,
                    human_score,
                    human_feedback,
                    created_at
                FROM {self.table_name}
                WHERE business_id = %s
                  AND content_type = %s 
                  AND (human_approved = TRUE OR (human_approved IS NULL AND agent_auto_score >= 8.0))
                ORDER BY 
                    CASE WHEN human_approved = TRUE THEN 1 ELSE 2 END,
                    COALESCE(human_score, agent_auto_score) DESC,
                    created_at DESC
                LIMIT %s
            """, (self.business_id, content_type, limit))
            
            results = cursor.fetchall()
            
            patterns = []
            for row in results:
                patterns.append({
                    'generation_id': row['generation_id'],
                    'creative_angle': row['creative_angle'],
                    'content': row['generated_content'],
                    'auto_score': float(row['auto_score']) if row['auto_score'] else 0.0,
                    'human_approved': row['human_approved'],
                    'human_score': float(row['human_score']) if row['human_score'] else None,
                    'human_feedback': row['human_feedback'],
                    'created_at': row['created_at'],
                    'pattern_data': {
                        'creative_angle': row['creative_angle'],
                        'what_worked': row['human_feedback'] if row['human_approved'] else ""
                    }
                })
            
            return patterns
            
        except Exception as e:
            logger.error(f"Failed to retrieve approved patterns: {e}", exc_info=True)
            return []
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((psycopg2.OperationalError, psycopg2.DatabaseError))
    )
    def get_rejected_patterns(self, content_type: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Get rejected patterns using raw SQL"""
        conn = None
        cursor = None
        
        try:
            conn = db_circuit_breaker.call(self._get_connection)
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cursor.execute(f"""
                SELECT 
                    generation_id,
                    creative_angle,
                    agent_auto_score as auto_score,
                    human_feedback,
                    created_at
                FROM {self.table_name}
                WHERE business_id = %s
                  AND content_type = %s 
                  AND (human_approved = FALSE OR (human_approved IS NULL AND agent_auto_score < 7.0))
                ORDER BY created_at DESC
                LIMIT %s
            """, (self.business_id, content_type, limit))
            
            results = cursor.fetchall()
            
            patterns = []
            for row in results:
                patterns.append({
                    'generation_id': row['generation_id'],
                    'creative_angle': row['creative_angle'],
                    'auto_score': float(row['auto_score']) if row['auto_score'] else 0.0,
                    'human_feedback': row['human_feedback'] or "Low alignment score",
                    'created_at': row['created_at'],
                    'pattern_data': {
                        'creative_angle': row['creative_angle'],
                        'issue': row['human_feedback'] or "Failed automatic scoring"
                    }
                })
            
            return patterns
            
        except Exception as e:
            logger.error(f"Failed to retrieve rejected patterns: {e}", exc_info=True)
            return []
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()
    
    def get_learning_summary(self, content_type: str) -> str:
        """Generate summary of learned patterns"""
        approved = self.get_approved_patterns(content_type, limit=5)
        rejected = self.get_rejected_patterns(content_type, limit=3)
        
        summary = f"\n=== LEARNED PATTERNS FOR {content_type.upper()} ===\n\n"
        
        if approved:
            summary += "SUCCESSFUL APPROACHES (Human-Approved or High Auto-Score):\n"
            for i, pattern in enumerate(approved, 1):
                score = pattern.get('human_score') or pattern.get('auto_score', 0)
                approval_type = "[Human]" if pattern.get('human_approved') else "[Auto]"
                
                summary += f"{i}. {approval_type} Score: {score:.1f}/10\n"
                summary += f"   Creative angle: {pattern.get('creative_angle', 'N/A')}\n"
                summary += f"   What worked: {pattern.get('human_feedback') or 'Good alignment'}\n\n"
        
        if rejected:
            summary += "\nAVOID THESE APPROACHES (Rejected):\n"
            for i, pattern in enumerate(rejected, 1):
                summary += f"{i}. Score: {pattern.get('auto_score', 0):.1f}/10\n"
                summary += f"   What failed: {pattern.get('creative_angle', 'N/A')}\n"
                summary += f"   Feedback: {pattern.get('human_feedback', 'Low alignment')}\n\n"
        
        if not approved and not rejected:
            summary += "No learning data yet. Will learn from first attempt.\n"
        
        return summary
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type((psycopg2.OperationalError, psycopg2.DatabaseError))
    )
    def get_learning_stats(self, content_type: str) -> Dict[str, Any]:
        """Get statistics about learning progress using raw SQL"""
        conn = None
        cursor = None
        
        try:
            conn = db_circuit_breaker.call(self._get_connection)
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cursor.execute(f"""
                SELECT 
                    COUNT(*) as total_generations,
                    COUNT(CASE WHEN has_human_feedback = TRUE THEN 1 END) as human_feedback_count,
                    COUNT(CASE WHEN human_approved = TRUE THEN 1 END) as approved_count,
                    AVG(agent_auto_score) as avg_auto_score,
                    AVG(CASE WHEN human_score IS NOT NULL THEN human_score END) as avg_human_score,
                    COUNT(CASE WHEN agent_correct = TRUE THEN 1 END) as agent_correct_count
                FROM {self.table_name}
                WHERE business_id = %s
                  AND content_type = %s
            """, (self.business_id, content_type))
            
            result = cursor.fetchone()
            
            if result:
                human_feedback_count = result['human_feedback_count'] or 0
                
                return {
                    'total_generations': result['total_generations'] or 0,
                    'human_feedback_count': human_feedback_count,
                    'approved_count': result['approved_count'] or 0,
                    'approval_rate': (result['approved_count'] / human_feedback_count * 100) if human_feedback_count else 0,
                    'avg_auto_score': float(result['avg_auto_score'] or 0),
                    'avg_human_score': float(result['avg_human_score'] or 0),
                    'agent_accuracy': (result['agent_correct_count'] / human_feedback_count * 100) if human_feedback_count else 0
                }
            
            return {}
            
        except Exception as e:
            logger.error(f"Failed to get learning stats: {e}", exc_info=True)
            return {}
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()


# ===== RAG SYSTEM WITH RETRY LOGIC =====

class Marketing_Rag_System:
    """
    RAG system with production-ready features:
    - Retry logic for database operations
    - Connection pooling
    - Graceful degradation
    - Thread-safe initialization
    """
    
    def __init__(self, data_path: str = "./data/marketing", business_id=None, content_type: str = "default"):
        self._embed_model_cache = None
        self._setup_lock = threading.Lock()
        
        if not os.path.isabs(data_path):
            data_path = os.path.abspath(data_path)
        
        self.data_path = data_path
        self.content_type = content_type
        clean_id = (business_id or 'default').replace("-", "_")
        self.table_name = f"rag_data_{clean_id}_{content_type}"
        self.business_id = business_id
        get_init_lock(business_id)

    def initialize_embedding_model(self):
        """Initialize embedding model with caching"""
        if self._embed_model_cache is not None:
            return self._embed_model_cache
        
        with self._setup_lock:
            if self._embed_model_cache is not None:
                return self._embed_model_cache
            
            try:
                logger.info("Setting up embedding model...")
                self.embed_model = FastEmbedEmbedding(
                    model_name="BAAI/bge-small-en-v1.5"
                )
                self._embed_model_cache = self.embed_model
                logger.info("Embedding model ready!")
                return self.embed_model
            except Exception as e:
                logger.error(f"Failed to initialize embedding model: {e}", exc_info=True)
                raise

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        reraise=True
    )
    def build_marketing_index(self):
        """
        Build marketing index with retry logic and proper error handling.
        
        Returns:
            VectorStoreIndex or None if failed
        """
        from llama_index.llms.groq import Groq
        from llama_index.core import Settings
        
        try:
            rag_llm = Groq(model="llama-3.3-70b-versatile", api_key=GROQ_API_KEY)
                    
            Settings.embed_model = self.initialize_embedding_model()
            Settings.llm = rag_llm 
            
            if not self.business_id:
                logger.warning("Business ID not given")

            lock = get_init_lock(self.business_id)
            
            with lock:
                # Choose parser based on content type
                if self.content_type == "blog":
                    parser = SemanticSplitterNodeParser.from_defaults(
                        buffer_size=1, 
                        embed_model=Settings.embed_model,
                        breakpoint_percentile_threshold=95
                    )
                elif self.content_type in ["social", "ad", "email"]:
                    parser = SimpleNodeParser.from_defaults(chunk_size=256, chunk_overlap=50)
                else:
                    parser = SimpleNodeParser.from_defaults(chunk_size=512, chunk_overlap=100)
                
                # Parse PostgreSQL URI
                from urllib.parse import urlparse
                raw_uri = os.getenv("POSTGRES_ASYNC_URI", "").strip()
                raw_uri = raw_uri.replace("&channel_binding=require", "")
                raw_uri = raw_uri.replace("channel_binding=require&", "")
                raw_uri = raw_uri.replace("?channel_binding=require", "")
                parsed = urlparse(raw_uri)
                
                logger.info("="*80)
                logger.info(f"Building index for: {self.business_id} / {self.content_type}")
                logger.info(f"Data path: {self.data_path}")
                logger.info(f"Table name: {self.table_name}")
                
                # Validate data path
                if not os.path.exists(self.data_path):
                    logger.error(f"Data path does not exist: {self.data_path}")
                    raise ValueError(f"Data path not found: {self.data_path}")
                
                files = os.listdir(self.data_path)
                logger.info(f"Files found: {len(files)}")
                for f in files[:5]:
                    logger.info(f"   - {f}")
                
                if not files:
                    logger.warning(f"No files in {self.data_path}")
                    return None
                
                logger.info("="*80)
                
                # Create vector store
                self.vector_store = PGVectorStore.from_params(
                    host=parsed.hostname,
                    port=parsed.port or 5432,
                    database=parsed.path.lstrip('/'),
                    user=parsed.username,
                    password=parsed.password,
                    table_name=self.table_name,
                    embed_dim=384,
                    hybrid_search=True,
                    hnsw_kwargs={
                        "hnsw_m": 16,
                        "hnsw_ef_construction": 64,
                        "hnsw_ef_search": 40,
                        "hnsw_dist_method": "vector_cosine_ops"
                    }
                )

                Settings.node_parser = parser
                storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
                
                # Try to load existing index
                try:
                    index = VectorStoreIndex.from_vector_store(
                        vector_store=self.vector_store,
                        embed_model=Settings.embed_model
                    )
                    logger.info(f"Loaded existing index: {self.table_name}")
                    
                    # Verify index
                    retriever = index.as_retriever(similarity_top_k=3)
                    test_result = retriever.retrieve("test query")
                    logger.info(f"Index verification: Found {len(test_result)} nodes")
                    
                    if len(test_result) == 0:
                        logger.warning("Index exists but is empty! Rebuilding...")
                        raise Exception("Empty index")
                        
                except Exception as load_error:
                    logger.info(f"Building NEW index for {self.table_name}")
                    
                    # PRIORITY 1: Load documents from DATABASE
                    documents = []
                    try:
                        from llama_index.core import Document
                        from db import BrandDocument
                        
                        logger.info(f"Checking database for documents for {self.business_id}/{self.content_type}...")
                        with get_db_session() as session:
                            db_docs = session.query(BrandDocument).filter_by(
                                business_id=self.business_id,
                                content_type=self.content_type
                            ).all()
                            
                            logger.info(f"Found {len(db_docs)} documents in database")
                            
                            for db_doc in db_docs:
                                if db_doc.file_content:
                                    # Create LlamaIndex Document from database content
                                    doc = Document(
                                        text=db_doc.file_content,
                                        metadata={
                                            "filename": db_doc.filename,
                                            "business_id": db_doc.business_id,
                                            "content_type": db_doc.content_type,
                                            "source": "database"
                                        }
                                    )
                                    documents.append(doc)
                                    logger.info(f"   ✓ Loaded from DB: {db_doc.filename}")

                    except Exception as db_error:
                        logger.warning(f"Database document loading failed: {db_error}")

                    # FALLBACK: Load from filesystem if database is empty or failed
                    if not documents:
                        logger.info("No documents in database, checking filesystem...")
                        if os.path.exists(self.data_path):
                            documents = SimpleDirectoryReader(self.data_path).load_data()
                    
                    if not documents:
                        logger.error(f"No documents loaded for {self.business_id}/{self.content_type}")
                        return None
                    
                    logger.info(f"Total documents for indexing: {len(documents)}")
                    
                    # Build index with BATCHING to avoid SSL/Connection errors
                    logger.info("Building vector index (Batched)...")
                    
                    BATCH_SIZE = 5
                    total_docs = len(documents)
                    
                    # Initialize with first batch
                    first_batch = documents[:BATCH_SIZE]
                    logger.info(f"Processing Initial Batch 1/{(total_docs + BATCH_SIZE - 1)//BATCH_SIZE} ({len(first_batch)} docs)")
                    
                    index = VectorStoreIndex.from_documents(
                        first_batch,
                        storage_context=storage_context,
                        show_progress=True
                    )
                    
                    # Insert remaining batches
                    if total_docs > BATCH_SIZE:
                        import time
                        for i in range(BATCH_SIZE, total_docs, BATCH_SIZE):
                            batch = documents[i : i + BATCH_SIZE]
                            batch_num = i//BATCH_SIZE + 1
                            total_batches = (total_docs + BATCH_SIZE - 1)//BATCH_SIZE
                            
                            logger.info(f"Processing Batch {batch_num}/{total_batches} ({len(batch)} docs)")
                            
                            try:
                                # Small delay to let DB breathe
                                time.sleep(0.5) 
                                for doc in batch:
                                    index.insert(doc)
                            except Exception as batch_e:
                                logger.error(f"Batch {batch_num} failed: {batch_e}")
                                raise batch_e
                    logger.info(f"Index built successfully: {self.table_name}")
                
                logger.info("="*80)
                return index
                
        except Exception as e:
            logger.error(f"Index build failed: {e}", exc_info=True)
            raise


# ===== KNOWLEDGE BASE WITH CACHING =====

VALID_CONTENT_TYPES = {"blog", "social", "ad"}

class BrandVoiceKnowledgeBase:
    """
    Knowledge base with caching and error handling.
    Thread-safe implementation.
    """
    
    def __init__(self, postgres_uri: Optional[str] = None, llm=None):
        self.postgres_uri = os.getenv("POSTGRES_ASYNC_URI") 
        self.llm = llm
        self._kb_cache = {}
        self._load_errors = {}
        self._cache_lock = Lock()

    def _get_kb(self, content_type: str, business_id: str):
        """Get knowledge base with caching and thread safety"""
        cache_key = (content_type, business_id)

        with self._cache_lock:
            if cache_key in self._kb_cache:
                logger.info(f"Using cached KB for {business_id}/{content_type}")
                return self._kb_cache[cache_key]
        
        lock = get_init_lock(f"{business_id}_{content_type}")

        with lock:
            # Double-check after acquiring lock
            with self._cache_lock:
                if cache_key in self._kb_cache:
                    return self._kb_cache[cache_key]
            
            try:
                data_path = os.path.abspath(f"brand_{content_type}s/{business_id}")
                
                logger.info(f"Creating new KB for {business_id}/{content_type}")
                
                rag_system = Marketing_Rag_System(
                    data_path=data_path,
                    content_type=content_type,
                    business_id=business_id
                )
                kb = rag_system.build_marketing_index()
                
                with self._cache_lock:
                    self._kb_cache[cache_key] = kb
                
                return kb
                
            except Exception as e:
                logger.error(f"Failed to load {content_type} KB: {e}", exc_info=True)
                self._load_errors[cache_key] = str(e)
                raise
    
    def _get_llm(self):
        """Get LLM with lazy initialization"""
        if self.llm is None:
            self.llm = get_llm()
        return self.llm
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        reraise=True
    )
    def get_style_guide(self, content_type: str, business_id: str):
        """
        Get style guide query engine with retry logic.
        
        Args:
            content_type: blog/social/ad
            business_id: Business identifier
            
        Returns:
            Query engine or raises exception
        """
        if content_type not in VALID_CONTENT_TYPES:
            raise ValueError(f"Invalid content_type. Must be one of: {VALID_CONTENT_TYPES}")
        
        self.business_id = business_id
        
        try:
            kb = self._get_kb(content_type, business_id)
            
            if kb is None:
                raise ValueError(f"KB for {content_type} returned None")
            
            top_k_map = {"blog": 5, "social": 3, "ad": 3}
            
            return kb.as_query_engine(
                similarity_top_k=top_k_map[content_type],
                response_mode="compact"
            )
            
        except Exception as e:
            logger.error(f"get_style_guide failed: {e}", exc_info=True)
            raise# ===== TOOLS WITH RATE LIMITING =====

class KBQueryTool(BaseTool):
    """Knowledge base query tool with error handling"""
    
    name: str = "kb_query"
    description: str = """Query brand knowledge base for style examples.
    
    IMPORTANT: content_type must be EXACTLY one of: 'blog', 'social', or 'ad'
    - Use 'blog' for blog articles, long-form content
    - Use 'social' for social media posts, tweets, LinkedIn posts
    - Use 'ad' for advertisements, ad copy
    
    Example: {"content_type": "blog", "query": "tone examples"}
    """
    kb: Any = Field(default=None, exclude=True)
    business_id: str = Field(default="default", exclude=True)
    
    def __init__(self, kb: BrandVoiceKnowledgeBase, business_id: str):
        super().__init__()
        object.__setattr__(self, 'kb', kb)
        object.__setattr__(self, 'business_id', business_id)
    
    @retry(
        stop=stop_after_attempt(2),
        wait=wait_exponential(multiplier=1, min=1, max=5),
        reraise=False
    )
    def _run(self, content_type: str, query: str) -> str:
        """Execute KB query with retry logic"""
        logger.info("="*80)
        logger.info(f"KB QUERY - {content_type}: {query}")
        logger.info("="*80)
        
        if not content_type or not query:
            return "Error: Missing content_type or query"
        
        if content_type not in VALID_CONTENT_TYPES:
            return f"Error: content_type must be one of {VALID_CONTENT_TYPES}"
        
        try:
            style_guide = self.kb.get_style_guide(content_type, self.business_id)
            response = style_guide.query(query)
            response_str = str(response)
            
            logger.info(f"KB Response: {len(response_str)} chars")
            
            if not response_str or response_str.strip() == "":
                logger.warning("Empty KB response")
                return self._get_fallback_response(content_type)
            
            return response_str
            
        except Exception as e:
            logger.error(f"KB query failed: {e}", exc_info=True)
            return self._get_fallback_response(content_type)
    
    def _get_fallback_response(self, content_type: str) -> str:
        """Fallback response when KB fails"""
        fallbacks = {
            "blog": "BRAND FALLBACK: Conversational yet professional tone, clear structure, practical insights",
            "social": "BRAND FALLBACK: Engaging, brief, personality-driven",
            "ad": "BRAND FALLBACK: Clear value proposition, compelling call-to-action"
        }
        return fallbacks.get(content_type, fallbacks["blog"])


class TavilySearchTool(BaseTool):
    """Web search tool with rate limiting and circuit breaker"""
    
    name: str = "tavily_search"
    description: str = "Search the web using Tavily API for current information and trends"
    
    @sleep_and_retry
    @tavily_rate_limit
    @tavily_circuit_breaker
    @retry(
        stop=stop_after_attempt(2),
        wait=wait_exponential(multiplier=1, min=2, max=5),
        reraise=False
    )
    def _run(self, query: str) -> str:
        """Execute search with rate limiting and retry logic"""
        if not TAVILY_API_KEY:
            return "Web search unavailable - TAVILY_API_KEY not configured"
        
        try:
            client = TavilyClient(api_key=TAVILY_API_KEY)
            result = client.search(query, max_results=5)
            return str(result)
            
        except Exception as e:
            logger.error(f"Search failed: {e}")
            return f"Search temporarily unavailable. Error: {str(e)[:100]}"


class LearningMemoryTool(BaseTool):
    """Tool to access learned brand patterns"""
    
    name: str = "learning_memory"
    description: str = """Access learned brand patterns from past human feedback and auto-scores.
    
    Actions:
    - get_approved: Get successful creative approaches (human or auto-approved)
    - get_rejected: Get patterns that were rejected (to avoid them)
    - get_summary: Get full learning summary for context
    """
    
    memory: Any = Field(default=None, exclude=True)
    
    def __init__(self, memory: BrandLearningMemory):
        super().__init__()
        object.__setattr__(self, 'memory', memory)
    
    @retry(
        stop=stop_after_attempt(2),
        wait=wait_exponential(multiplier=1, min=1, max=5),
        reraise=False
    )
    def _run(self, action: str, content_type: str = "blog") -> str:
        """Access learning memory with error handling"""
        try:
            if action == "get_approved":
                patterns = self.memory.get_approved_patterns(content_type, limit=10)
                if not patterns:
                    return "No approved patterns yet. This is your first attempt!"
                
                result = "SUCCESSFUL CREATIVE APPROACHES:\n"
                for p in patterns:
                    data = p.get('pattern_data', {})
                    if isinstance(data, str):
                        try:
                            data = json.loads(data)
                        except:
                            data = {}
                    
                    approval_type = "[Human]" if p.get('human_approved') else "[Auto]"
                    score = p.get('human_score') or p.get('auto_score', 0)
                    result += f"- {approval_type} | {data.get('creative_angle', 'N/A')} (Score: {score:.1f}/10)\n"
                return result
            
            elif action == "get_rejected":
                patterns = self.memory.get_rejected_patterns(content_type, limit=5)
                if not patterns:
                    return "No rejected patterns yet."
                
                result = "AVOID THESE APPROACHES:\n"
                for p in patterns:
                    data = p.get('pattern_data', {})
                    if isinstance(data, str):
                        try:
                            data = json.loads(data)
                        except:
                            data = {}
                    result += f"- {data.get('issue', 'N/A')}: {p.get('human_feedback', 'Low score')}\n"
                return result
            
            elif action == "get_summary":
                return self.memory.get_learning_summary(content_type)
            
            else:
                return f"Unknown action: {action}"
                
        except Exception as e:
            logger.error(f"Learning memory access failed: {e}", exc_info=True)
            return f"Memory temporarily unavailable. Error: {str(e)[:100]}"


class BrandMetricsTool(BaseTool):
    """Access concrete brand metrics"""
    
    name: str = "brand_metrics"
    description: str = """Get concrete brand metrics extracted from documents.
    
    Returns: target sentence length, paragraph structure, signature phrases, etc.
    Use this to compare your content against measurable brand standards.
    """
    
    analyzer: Any = Field(default=None, exclude=True)
    
    def __init__(self, analyzer: BrandMetricsAnalyzer):
        super().__init__()
        object.__setattr__(self, 'analyzer', analyzer)
    
    def _run(self, action: str = "get_metrics") -> str:
        """Get brand metrics"""
        try:
            if action == "get_metrics":
                metrics = self.analyzer.metrics
                
                result = "BRAND METRICS (GROUND TRUTH):\n\n"
                result += f"Target Sentence Length: {metrics['target_sentence_length']:.1f} words\n"
                result += f"Target Sentences/Paragraph: {metrics['target_sentences_per_para']:.1f}\n"
                result += f"Uses Bullet Points: {metrics['uses_bullets']}\n"
                result += f"Uses Numbered Lists: {metrics['uses_numbered_lists']}\n"
                result += f"\nSignature Phrases ({len(metrics['signature_phrases'])}):\n"
                for i, phrase in enumerate(metrics['signature_phrases'][:15], 1):
                    result += f"{i}. {phrase}\n"
                
                return result
            
            else:
                return f"Unknown action: {action}"
                
        except Exception as e:
            logger.error(f"Metrics access failed: {e}", exc_info=True)
            return f"Metrics temporarily unavailable. Error: {str(e)[:100]}"


class ContentScoringTool(BaseTool):
    """Automatically score content against brand metrics"""
    
    name: str = "score_content"
    description: str = """Automatically score content structure against brand metrics.
    
    Pass the full generated content as a string.
    Returns objective scores based on actual measurements.
    """
    
    analyzer: Any = Field(default=None, exclude=True)
    
    def __init__(self, analyzer: BrandMetricsAnalyzer):
        super().__init__()
        object.__setattr__(self, 'analyzer', analyzer)
    
    def _run(self, content: str) -> str:
        """Score content with error handling"""
        try:
            scores = self.analyzer.score_content(content)
            
            result = "AUTOMATIC STRUCTURAL SCORES (Objective Measurements):\n\n"
            result += f"Overall Structure Score: {scores['structure_score']}/10\n"
            result += f"Signature Phrase Score: {scores['phrase_score']}/10\n\n"
            
            result += "DETAILED BREAKDOWN:\n"
            result += f"- Sentence Length Score: {scores['sentence_length_score']}/10\n"
            result += f"  Target: {scores['details']['target_sentence_length']} words\n"
            result += f"  Actual: {scores['details']['actual_sentence_length']} words\n\n"
            
            result += f"- Paragraph Structure Score: {scores['para_structure_score']}/10\n"
            result += f"  Target: {scores['details']['target_sentences_per_para']} sentences/para\n"
            result += f"  Actual: {scores['details']['actual_sentences_per_para']} sentences/para\n\n"
            
            result += f"- Format Alignment Score: {scores['format_score']}/10\n\n"
            
            result += f"Signature Phrases:\n"
            result += f"- Available: {scores['details']['signature_phrases_available']}\n"
            result += f"- Used: {scores['details']['signature_phrases_used']}\n"
            phrases_str = ', '.join(scores['details']['phrases_found']) if scores['details']['phrases_found'] else 'None'
            result += f"- Found: {phrases_str}\n"
            
            return result
            
        except Exception as e:
            logger.error(f"Content scoring failed: {e}", exc_info=True)
            return f"Scoring temporarily unavailable. Error: {str(e)[:100]}"


# Global tool instances
tavily_search = TavilySearchTool()


# ===== NEW HELPER FUNCTIONS FOR ITERATIVE GENERATION =====

def extract_iteration_feedback(critic_output: str) -> Dict[str, Any]:
    """
    Parse critic's structured feedback to extract actionable issues.
    
    Args:
        critic_output: Raw output from critic agent
        
    Returns:
        Dict with parsed feedback including issues and recommendations
    """
    try:
        # Try to find JSON in the output
        json_match = re.search(r'\{.*?"decision".*?\}', critic_output, re.DOTALL)
        
        if json_match:
            feedback_data = json.loads(json_match.group(0))
            return {
                'decision': feedback_data.get('decision', 'revise'),
                'current_score': feedback_data.get('current_score', 0.0),
                'issues': feedback_data.get('issues', []),
                'specific_fixes': feedback_data.get('specific_fixes', []),
                'what_to_keep': feedback_data.get('what_to_keep', [])
            }
        
        # Fallback: extract issues from text
        issues = []
        if 'sentence length' in critic_output.lower():
            issues.append("Adjust sentence length to match target")
        if 'phrase' in critic_output.lower() and 'missing' in critic_output.lower():
            issues.append("Incorporate more signature phrases")
        if 'paragraph' in critic_output.lower():
            issues.append("Adjust paragraph structure")
        
        return {
            'decision': 'revise' if issues else 'approve',
            'current_score': 7.0,
            'issues': issues,
            'specific_fixes': issues,
            'what_to_keep': []
        }
        
    except Exception as e:
        logger.error(f"Failed to parse critic feedback: {e}")
        return {
            'decision': 'revise',
            'current_score': 7.0,
            'issues': ['General improvement needed'],
            'specific_fixes': ['Review and refine content'],
            'what_to_keep': []
        }


def format_iteration_context(previous_content: str, feedback: Dict[str, Any], iteration_num: int) -> str:
    """
    Format previous attempt and feedback for next iteration.
    
    Args:
        previous_content: The content from previous iteration
        feedback: Parsed feedback from critic
        iteration_num: Current iteration number
        
    Returns:
        Formatted context string for writer agent
    """
    context = f"""
=== ITERATION {iteration_num} REFINEMENT ===

PREVIOUS ATTEMPT:
{previous_content}

CRITIC FEEDBACK (Score: {feedback.get('current_score', 0)}/10):

ISSUES TO FIX:
"""
    for i, issue in enumerate(feedback.get('issues', []), 1):
        context += f"{i}. {issue}\n"
    
    if feedback.get('specific_fixes'):
        context += "\nSPECIFIC ACTIONS:\n"
        for i, fix in enumerate(feedback.get('specific_fixes', []), 1):
            context += f"{i}. {fix}\n"
    
    if feedback.get('what_to_keep'):
        context += "\nKEEP THESE ELEMENTS (they're working well):\n"
        for i, element in enumerate(feedback.get('what_to_keep', []), 1):
            context += f"{i}. {element}\n"
    
    context += """
YOUR TASK: Revise ONLY the problematic areas. Keep what's working.
Do NOT start from scratch - build on the previous attempt.
"""
    
    return context


def extract_content_from_writer_output(output: str) -> str:
    """
    Extract clean content from writer's output, removing metadata.
    
    Args:
        output: Raw output from writer agent
        
    Returns:
        Clean content string
    """
    try:
        # Look for [FINAL CONTENT] tags
        content_match = re.search(r'\[FINAL CONTENT\](.*?)\[/FINAL CONTENT\]', output, re.DOTALL | re.IGNORECASE)
        
        if content_match:
            content = content_match.group(1).strip()
            logger.info(f"Extracted content from [FINAL CONTENT] tags: {len(content)} chars")
            return content
        
        # Fallback: remove JSON blocks and metadata
        content = output
        content = re.sub(r'\{[^}]*"decision"[^}]*\}', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'\{[^}]*"final_score"[^}]*\}', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'\[METADATA\].*?\[/METADATA\]', '', content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove iteration logs
        content = re.sub(r'\[ITERATION LOG\].*?\[/ITERATION LOG\]', '', content, flags=re.DOTALL | re.IGNORECASE)
        
        content = content.strip()
        logger.info(f"Extracted content via fallback: {len(content)} chars")
        return content
        
    except Exception as e:
        logger.error(f"Content extraction failed: {e}")
        return output.strip()
    

# ===== ITERATIVE GENERATION WITH CRITIC LOOP =====

@retry(
    stop=stop_after_attempt(2),
    wait=wait_exponential(multiplier=2, min=4, max=20),
    reraise=True
)
def run_iterative_generation(business_id: str, topic: str, 
                             format_type: str, voice: str) -> Tuple[Dict[str, Any], str]:
    """
    Generate content using iterative refinement with external critic feedback.
    
    Process:
    1. Initial generation by writer
    2. Critic reviews and provides feedback
    3. Writer revises based on specific feedback
    4. Repeat until score >= 8.0 OR max 3 iterations
    
    Returns:
        (result_dict, generation_id)
    """
    
    generation_id = str(uuid.uuid4())
    
    # Get user_id
    user_id = None
    try:
        with get_db_session() as session:
            user = session.query(User).filter_by(business_id=business_id).first()
            if user:
                user_id = user.id
            else:
                user = User(
                    email=f"{business_id}@brandguard.ai",
                    business_id=business_id,
                    business_name=business_id,
                    industry='Marketing'
                )
                session.add(user)
                session.commit()
                session.refresh(user)
                user_id = user.id
                logger.info(f"Created new user for {business_id}: user_id={user_id}")
    except Exception as e:
        logger.error(f"Failed to get user_id: {e}", exc_info=True)
    
    logger.info(f"Generation ID: {generation_id}, User ID: {user_id}")
    
    # Determine content type
    if "blog" in format_type.lower():
        content_type = "blog"
    elif "social" in format_type.lower():
        content_type = "social"
    elif "ad" in format_type.lower():
        content_type = "ad"
    else:
        content_type = "blog"
    
    # Initialize components
    kb = BrandVoiceKnowledgeBase()
    learning_memory = BrandLearningMemory(business_id=business_id)
    metrics_analyzer = BrandMetricsAnalyzer(business_id=business_id, content_type=content_type)
    
    factory = ContentCrewFactory(
        kb=kb,
        tavily_search=tavily_search,
        learning_memory=learning_memory,
        metrics_analyzer=metrics_analyzer,
        business_id=business_id
    )
    
    # Create base crew (research, strategy, brand analysis)
    base_crew = factory.create_base_crew()
    
    # Run base crew to get context
    logger.info(f"Running base crew (research, strategy, brand analysis)...")
    base_inputs = {
        "topic": topic,
        "format": format_type,
        "voice": voice,
        "content_type": content_type
    }
    base_result = base_crew.kickoff(inputs=base_inputs)
    
    logger.info(f"Base crew complete. Starting iterative refinement...")
    
    # Iterative refinement loop
    MAX_ITERATIONS = 3
    TARGET_SCORE = 8.0
    
    current_content = None
    current_score = 0.0
    creative_angle = "Unknown"
    iteration_history = []
    
    for iteration in range(1, MAX_ITERATIONS + 1):
        logger.info(f"=== ITERATION {iteration}/{MAX_ITERATIONS} ===")
        
        # Create writer crew for this iteration
        if iteration == 1:
            # First iteration: write from scratch
            writer_crew = factory.create_writer_crew(
                iteration_num=iteration,
                previous_content=None,
                feedback=None
            )
            writer_inputs = {**base_inputs, "base_context": str(base_result)}
        else:
            # Subsequent iterations: revise based on feedback
            writer_crew = factory.create_writer_crew(
                iteration_num=iteration,
                previous_content=current_content,
                feedback=feedback
            )
            writer_inputs = {
                **base_inputs,
                "base_context": str(base_result),
                "previous_content": current_content,
                "feedback": format_iteration_context(current_content, feedback, iteration)
            }
        
        # Run writer
        writer_result = writer_crew.kickoff(inputs=writer_inputs)
        current_content = extract_content_from_writer_output(str(writer_result))
        
        if not current_content or len(current_content.strip()) < 50:
            logger.error(f"Iteration {iteration}: Content too short or empty!")
            current_content = str(writer_result)
        
        logger.info(f"Iteration {iteration}: Generated {len(current_content)} chars")
        
        # Create critic crew to evaluate
        critic_crew = factory.create_critic_crew()
        critic_inputs = {
            **base_inputs,
            "content_to_review": current_content
        }
        
        critic_result = critic_crew.kickoff(inputs=critic_inputs)
        
        # Parse critic feedback
        feedback = extract_iteration_feedback(str(critic_result))
        current_score = feedback.get('current_score', 0.0)
        decision = feedback.get('decision', 'revise')
        
        # Extract creative angle if available
        angle_match = re.search(r'"creative_angle"\s*:\s*"([^"]+)"', str(writer_result))
        if angle_match:
            creative_angle = angle_match.group(1)
        
        iteration_history.append({
            'iteration': iteration,
            'score': current_score,
            'decision': decision,
            'issues': feedback.get('issues', []),
            'content_length': len(current_content)
        })
        
        logger.info(f"Iteration {iteration}: Score={current_score:.1f}, Decision={decision}")
        
        # Check if we should stop
        if decision == 'approve' or current_score >= TARGET_SCORE:
            logger.info(f"✓ Quality threshold reached at iteration {iteration}")
            break
        
        if iteration == MAX_ITERATIONS:
            logger.info(f"⚠ Reached max iterations ({MAX_ITERATIONS})")
            break
    
    # Final validation
    logger.info("Running final validator...")
    validator_crew = factory.create_validator_crew()
    validator_inputs = {
        **base_inputs,
        "final_content": current_content,
        "iteration_count": len(iteration_history),
        "final_score_from_critic": current_score
    }
    
    validator_result = validator_crew.kickoff(inputs=validator_inputs)
    
    # Parse final result
    try:
        json_match = re.search(r'\{.*?"validation_status".*?\}', str(validator_result), re.DOTALL)
        if json_match:
            final_data = json.loads(json_match.group(0))
            final_score = final_data.get('final_score', current_score)
            validated_content = final_data.get('generated_content', current_content)
            
            # Use validated content if it's substantial
            if validated_content and len(validated_content.strip()) > 50:
                current_content = validated_content
        else:
            final_score = current_score
    except Exception as e:
        logger.error(f"Failed to parse validator output: {e}")
        final_score = current_score
    
    logger.info(f"Final score: {final_score:.1f}, Iterations: {len(iteration_history)}")
    
    # Save learning
    try:
        learning_memory.save_learning(
            generation_id=generation_id,
            content_type=content_type,
            creative_angle=creative_angle,
            generated_content=current_content,
            auto_score=final_score,
            topic=topic,
            format_type=format_type,
            user_id=user_id
        )
        logger.info(f"Saved learning for generation {generation_id}")
    except Exception as e:
        logger.error(f"Failed to save learning: {e}", exc_info=True)
    
    # Build result
    result_dict = {
        "content": current_content,
        "auto_score": final_score,
        "creative_angle": creative_angle,
        "generation_id": generation_id,
        "topic": topic,
        "format_type": format_type,
        "iteration_count": len(iteration_history),
        "iteration_history": iteration_history,
        "full_result": str(validator_result)
    }
    
    logger.info(f"Generation complete: {generation_id}")
    logger.info(f"  - Content: {len(current_content)} chars")
    logger.info(f"  - Score: {final_score}/10")
    logger.info(f"  - Iterations: {len(iteration_history)}")
    logger.info(f"  - Angle: {creative_angle}")
    
    return result_dict, generation_id


# ===== CONTENT CREW FACTORY WITH ITERATIVE DESIGN =====

class ContentCrewFactory:
    """
    Factory for creating specialized crews for iterative content generation.
    
    Crews:
    1. Base Crew: Research + Strategy + Brand Analysis (run once)
    2. Writer Crew: Generate or revise content (run per iteration)
    3. Critic Crew: Evaluate and provide feedback (run per iteration)
    4. Validator Crew: Final quality check (run once)
    """
    
    def __init__(self, kb: BrandVoiceKnowledgeBase, tavily_search: TavilySearchTool,
                 learning_memory: BrandLearningMemory, metrics_analyzer: BrandMetricsAnalyzer,
                 business_id: str):
        self.kb = kb
        self.tavily_search = tavily_search
        self.learning_memory = learning_memory
        self.metrics_analyzer = metrics_analyzer
        self.business_id = business_id
        
        # Get few-shot examples
        self.few_shot_examples = metrics_analyzer.get_top_examples(limit=3)
        if not self.few_shot_examples:
            self.few_shot_examples = "No previous examples available. Follow brand guidelines strictly."
        
        # Different LLMs for different agent types
        self.creative_llm = get_llm(temperature=0.8)
        self.analytical_llm = get_llm(temperature=0.5)
        
        # Initialize tools
        self.kb_tool = KBQueryTool(kb=self.kb, business_id=self.business_id)
        self.learning_tool = LearningMemoryTool(memory=self.learning_memory)
        self.metrics_tool = BrandMetricsTool(analyzer=self.metrics_analyzer)
        self.scoring_tool = ContentScoringTool(analyzer=self.metrics_analyzer)
    
    def create_base_crew(self) -> Crew:
        """
        Create base crew for research, strategy, and brand analysis.
        This runs ONCE per generation.
        """
        
        # Researcher Agent
        researcher_agent = Agent(
            role="Research Analyst",
            goal="Find current, factual information and diverse perspectives on {topic}",
            backstory="""You gather comprehensive research including:
                - Current trends and data
                - Multiple perspectives and angles
                - Unexpected insights and fresh takes
                - Statistics and expert opinions

                You provide factual information WITHOUT brand voice - that's the writer's job.""",
            tools=[self.tavily_search],
            llm=self.analytical_llm,
            verbose=True,
        )
        
        research_task = Task(
            description="""Research {topic} for {format} content.

            Find:
            1. Latest trends and developments (2026)
            2. Different angles and perspectives
            3. Surprising insights or contrarian views
            4. Data, statistics, expert quotes
            5. What competitors/others are saying

            Provide 10-15 diverse factual findings.""",
            expected_output="Comprehensive research with multiple angles and perspectives",
            agent=researcher_agent,
        )
        
        # Creative Strategist Agent
        creative_strategist = Agent(
            role="Creative Content Strategist",
            goal="Find unique, engaging angles that match brand voice but stand out",
            backstory="""You're a creative strategist who makes content memorable.

                    CRITICAL BALANCE:
                    - Stay true to brand voice (tone, values, vocabulary)
                    - Find fresh perspectives competitors haven't used
                    - Make it scroll-stopping while staying authentic

                    You review past learnings and propose 2-3 creative angles that are:
                    - Fresh and haven't been overused
                    - Aligned with brand personality
                    - Engaging and memorable
                    - Practically valuable""",
            tools=[self.kb_tool, self.learning_tool],
            llm=self.creative_llm,
            verbose=True,
        )
        
        creative_strategy_task = Task(
            description="""Develop creative content strategy for {topic} in {format}.

                    STEP 1 - Learn from past (MANDATORY):
                    Use learning_memory tool:
                    - {{"action": "get_rejected", "content_type": "{content_type}"}}
                    - {{"action": "get_approved", "content_type": "{content_type}"}}

                    Review:
                    - What creative angles WORKED (high scores, approved)?
                    - What creative angles FAILED (rejected, low scores)?

                    STEP 2 - Understand brand:
                    Query KB 2 times:
                    - {{"content_type": "{content_type}", "query": "brand personality and values"}}
                    - {{"content_type": "{content_type}", "query": "tone and voice characteristics"}}

                    STEP 3 - Propose 2-3 creative angles:

                    CRITICAL RULES:
                    - DO NOT propose angles similar to rejected patterns
                    - DO build on successful patterns
                    - If an angle was rejected before with feedback "too generic", 
                      ensure your new angles are MORE specific

                    ANGLE 1: [Name]
                    - Hook: [unique perspective]
                    - Why it's fresh: [differentiation]
                    - Past learning: [how it differs from rejected patterns]
                    - Example opening: "[sample]"

                    RECOMMENDED ANGLE: [Which one and why]""",
            agent=creative_strategist,
            context=[research_task],
            expected_output="2-3 creative angles with recommended approach"
        )
        
        # Brand Voice Analyst Agent
        brand_analyst = Agent(
            role="Brand Voice Pattern Analyst",
            goal="Extract concrete metrics and qualitative patterns from brand documents",
            backstory="""You analyze brand voice using:

                1. METRICS TOOL: Exact measurements (sentence length, paragraph structure, phrases)
                2. KB QUERIES: Qualitative patterns (tone, style, approach)

                You provide MEASURABLE CONSTRAINTS for the writer.""",
            tools=[self.kb_tool, self.metrics_tool],
            llm=self.analytical_llm,
            verbose=True,
        )
        
        brand_analysis_task = Task(
            description="""Extract brand voice patterns for {format}.

                STEP 1 - Get Metrics:
                Use brand_metrics tool: {{"action": "get_metrics"}}

                STEP 2 - Query KB (4 times):
                1. {{"content_type": "{content_type}", "query": "tone characteristics"}}
                2. {{"content_type": "{content_type}", "query": "opening and closing patterns"}}
                3. {{"content_type": "{content_type}", "query": "vocabulary and perspective"}}
                4. {{"content_type": "{content_type}", "query": "content structure and format rules"}}

                OUTPUT:

                STRUCTURAL METRICS:
                - Target sentence length: [X] words (±2 allowed)
                - Target sentences/paragraph: [X] (±1 allowed)
                - Format: [blog/email/social structure]
                - Signature phrases: [list top 10]

                TONE & STYLE:
                - Primary tone: [descriptors]
                - Opening pattern: [how to start]
                - Closing pattern: [how to end]
                - Perspective: [you/we/they]

                CRITICAL FORMAT RULES:
                - Content type: {format}
                - Structural requirements: [specific to type]
                - What to avoid: [format mistakes]""",
            agent=brand_analyst,
            context=[research_task],
            expected_output="Measurable brand blueprint with exact metrics and patterns"
        )
        
        return Crew(
            agents=[researcher_agent, creative_strategist, brand_analyst],
            tasks=[research_task, creative_strategy_task, brand_analysis_task],
            process=Process.sequential,
            verbose=True,
        )
    
    def create_writer_crew(self, iteration_num: int, previous_content: Optional[str] = None, 
                          feedback: Optional[Dict[str, Any]] = None) -> Crew:
        """
        Create writer crew for generating or revising content.
        
        Args:
            iteration_num: Current iteration number
            previous_content: Content from previous iteration (None for first iteration)
            feedback: Feedback from critic (None for first iteration)
        """
        
        if iteration_num == 1:
            # First iteration: generate from scratch
            writer_backstory = f"""You are a Brand Voice Mimic.
            
            **FEW-SHOT TRAINING (MIMIC THESE SAMPLES):**
            {self.few_shot_examples}
            
            Your job: Write content that perfectly matches the brand blueprint.
            Use the creative strategy and brand analysis provided to you.
            """
            
            writer_description = """Write {topic} content in {format} format.

            **PROCESS:**
            1. Review the brand blueprint (from brand analysis)
            2. Review the creative strategy (recommended angle)
            3. Write content that:
               - Matches target sentence length (±2 words)
               - Uses 3-5 signature phrases naturally
               - Follows the creative angle
               - Matches brand tone and perspective
            
            **OUTPUT:**
            
            [FINAL CONTENT]
            <Your content - no metadata here>
            [/FINAL CONTENT]
            
            [METADATA]
            {
                "creative_angle": "Brief description",
                "iteration": 1
            }
            [/METADATA]
            """
        else:
            # Subsequent iterations: revise based on feedback
            writer_backstory = f"""You are a Brand Voice Mimic with revision capabilities.
            
            **FEW-SHOT TRAINING:**
            {self.few_shot_examples}
            
            You've received specific feedback on your previous attempt.
            Your job: Fix ONLY the identified issues while keeping what works.
            """
            
            writer_description = """Revise your previous content based on critic feedback.

            **PREVIOUS ATTEMPT:**
            {previous_content}
            
            **FEEDBACK TO ADDRESS:**
            {feedback}
            
            **REVISION RULES:**
            1. Fix ONLY the specific issues mentioned
            2. Keep elements that were marked as "working well"
            3. Do NOT start from scratch
            4. Maintain the creative angle
            
            **OUTPUT:**
            
            [FINAL CONTENT]
            <Your revised content - no metadata here>
            [/FINAL CONTENT]
            
            [METADATA]
            {
                "creative_angle": "Same as before",
                "iteration": """ + str(iteration_num) + """
            }
            [/METADATA]
            """
        
        writer_agent = Agent(
            role="Brand Voice Writer",
            goal="Write or revise content to match brand metrics and creative strategy",
            backstory=writer_backstory,
            tools=[self.kb_tool, self.metrics_tool],
            llm=self.creative_llm,
            verbose=True,
        )
        
        writer_task = Task(
            description=writer_description,
            agent=writer_agent,
            expected_output="Content with metadata"
        )
        
        return Crew(
            agents=[writer_agent],
            tasks=[writer_task],
            process=Process.sequential,
            verbose=True,
        )
    
    def create_critic_crew(self) -> Crew:
        """
        Create critic crew to evaluate content and provide feedback.
        """
        
        critic_agent = Agent(
            role="Content Quality Critic",
            goal="Evaluate content against brand metrics and provide specific, actionable feedback",
            backstory="""You are an objective evaluator who:
            
            1. Measures content against concrete metrics
            2. Identifies specific issues (not vague feedback)
            3. Recommends targeted fixes
            4. Notes what's working well (to preserve it)
            
            You use the scoring tool to get objective measurements.
            """,
            tools=[self.scoring_tool, self.metrics_tool],
            llm=self.analytical_llm,
            verbose=True,
        )
        
        critic_task = Task(
            description="""Evaluate the content and provide structured feedback.

            **CONTENT TO REVIEW:**
            {content_to_review}
            
            **EVALUATION PROCESS:**
            
            1. **Score the content:**
               Call score_content tool: {{"content": "<full content here>"}}
            
            2. **Analyze the scores:**
               - Structure Score: Is it >= 8.0?
               - Phrase Score: Are signature phrases used?
               - Sentence/Paragraph: How far from target?
            
            3. **Identify specific issues:**
               Example issues:
               - "Sentences average 22 words, target is 15 → too long"
               - "Only 1 signature phrase used, need 3-5"
               - "Paragraphs have 6 sentences, target is 3-4"
            
            4. **Note what's working:**
               - "Creative angle is engaging"
               - "Tone matches brand"
               - "Opening is strong"
            
            **OUTPUT (JSON format):**
            
            {
                "decision": "approve" or "revise",
                "current_score": X.X,
                "issues": [
                    "Specific issue 1",
                    "Specific issue 2"
                ],
                "specific_fixes": [
                    "Break long sentences in paragraphs 2 and 4",
                    "Add signature phrase 'X' in section Y"
                ],
                "what_to_keep": [
                    "Strong opening hook",
                    "Good use of examples"
                ]
            }
            
            **DECISION CRITERIA:**
            - Approve: structure_score >= 8.0 AND phrase_score >= 6.0
            - Revise: anything below that
            """,
            agent=critic_agent,
            expected_output="JSON feedback with decision and specific issues"
        )
        
        return Crew(
            agents=[critic_agent],
            tasks=[critic_task],
            process=Process.sequential,
            verbose=True,
        )
    
    def create_validator_crew(self) -> Crew:
        """
        Create validator crew for final quality check.
        """
        
        validator_agent = Agent(
            role="Final Quality Validator",
            goal="Provide final confidence score and quality report",
            backstory="""You are the final checkpoint.
            
            The content has been through iterative refinement.
            Your job: Verify quality and provide confidence score.
            """,
            tools=[self.scoring_tool],
            llm=self.analytical_llm,
            verbose=True,
        )
        
        validator_task = Task(
            description="""Validate the final content.

            **FINAL CONTENT:**
            {final_content}
            
            **ITERATION INFO:**
            - Total iterations: {iteration_count}
            - Last score from critic: {final_score_from_critic}
            
            **VALIDATION:**
            
            1. **Independent verification:**
               Call score_content tool on the final content
            
            2. **Calculate confidence:**
               confidence = (verified_structure_score × 0.6) + (phrase_score × 0.4)
            
            3. **Output assessment:**
            
            {
                "validation_status": "PASSED",
                "verified_structure_score": X.X,
                "verified_phrase_score": X.X,
                "confidence_score": X.X,
                "iteration_count": X,
                "creative_angle": "Extract from content context",
                "generated_content": "<EXACT copy of final_content>",
                "final_score": X.X
            }
            
            **CRITICAL:** Copy the final_content EXACTLY. Do not modify.
            """,
            agent=validator_agent,
            expected_output="JSON validation report with final content"
        )
        
        return Crew(
            agents=[validator_agent],
            tasks=[validator_task],
            process=Process.sequential,
            verbose=True,
        )# ===== UPDATE GENERATION WITH HUMAN FEEDBACK =====

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=10),
    reraise=True
)
def update_generation_with_human_feedback(generation_id: str, 
                                         business_id: str,
                                         human_approved: bool, 
                                         human_score: float,
                                         human_feedback: str) -> bool:
    """Update existing learning with human feedback"""
    learning_memory = BrandLearningMemory(business_id=business_id)
    
    success = learning_memory.update_with_human_feedback(
        generation_id=generation_id,
        human_approved=human_approved,
        human_score=human_score,
        human_feedback=human_feedback
    )
    
    if success:
        content_type = "blog"
        stats = learning_memory.get_learning_stats(content_type)
        
        logger.info(f"Human feedback saved for {generation_id}")
        logger.info(f"Agent Accuracy: {stats.get('agent_accuracy', 0):.1f}%")
        logger.info(f"Approval Rate: {stats.get('approval_rate', 0):.1f}%")
    else:
        logger.warning(f"Generation {generation_id} not found")
    
    return success


# ===== BACKWARD COMPATIBILITY WRAPPER =====

def run_generation_with_learning(business_id: str, topic: str, 
                                 format_type: str, voice: str) -> Tuple[Dict[str, Any], str]:
    """
    Backward compatibility wrapper - calls run_iterative_generation.
    
    This ensures existing code that calls run_generation_with_learning still works.
    """
    logger.info("run_generation_with_learning called - routing to run_iterative_generation")
    return run_iterative_generation(business_id, topic, format_type, voice)


# ===== VERIFICATION AND HELPER FUNCTIONS =====

def verify_learning_loop(business_id: str, content_type: str = "blog"):
    """Debug function to check if learning system is working"""
    memory = BrandLearningMemory(business_id=business_id)
    
    stats = memory.get_learning_stats(content_type)
    
    print("\n" + "="*60)
    print("LEARNING SYSTEM STATUS")
    print("="*60)
    print(f"Business ID: {business_id}")
    print(f"Content Type: {content_type}")
    print("-"*60)
    print(f"Total Generations: {stats.get('total_generations', 0)}")
    print(f"Human Feedback Count: {stats.get('human_feedback_count', 0)}")
    print(f"Approval Rate: {stats.get('approval_rate', 0):.1f}%")
    print(f"Avg Auto Score: {stats.get('avg_auto_score', 0):.1f}/10")
    print(f"Avg Human Score: {stats.get('avg_human_score', 0):.1f}/10")
    print(f"Agent Accuracy: {stats.get('agent_accuracy', 0):.1f}%")
    print("="*60)
    
    if stats.get('total_generations', 0) == 0:
        print("WARNING: No learnings saved!")
    elif stats.get('agent_accuracy', 0) < 50 and stats.get('human_feedback_count', 0) > 5:
        print("WARNING: Agent accuracy low!")
    else:
        print("Learning system operational")
    
    return stats


# ===== FACTORY FUNCTION =====

def get_crew(business_id: str = "default", topic: str = "AI Agents", 
             format_type: str = "Blog Article", voice: str = "formal",
             previous_feedback: str = None) -> tuple:
    """
    Create a crew with all components.
    
    NOTE: This is for backward compatibility. The new iterative system
    creates crews dynamically per iteration.
    
    Returns: (crew, kb, learning_memory, metrics_analyzer)
    """
    logger.info("get_crew called - NOTE: Iterative system uses dynamic crew creation", 
                business_id=business_id, 
                topic=topic, 
                format=format_type,
                voice=voice,
                has_feedback=bool(previous_feedback))

    load_dotenv()
    
    # Determine content type
    if "blog" in format_type.lower():
        content_type = "blog"
    elif "social" in format_type.lower():
        content_type = "social"
    elif "ad" in format_type.lower() or "advertisement" in format_type.lower():
        content_type = "ad"
    else:
        content_type = "blog"
    
    kb = BrandVoiceKnowledgeBase()
    learning_memory = BrandLearningMemory(business_id=business_id)
    metrics_analyzer = BrandMetricsAnalyzer(business_id=business_id, content_type=content_type)
    
    factory = ContentCrewFactory(
        kb=kb,
        tavily_search=tavily_search,
        learning_memory=learning_memory,
        metrics_analyzer=metrics_analyzer,
        business_id=business_id
    )
    
    # Return base crew for compatibility
    crew = factory.create_base_crew()
    return crew, kb, learning_memory, metrics_analyzer


# ===== CLI ENTRY POINT =====

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate brand-aligned marketing content with iterative refinement")
    
    parser.add_argument("--topic", default="AI Agents", help="Content topic")
    parser.add_argument("--format", default="Blog Article", help="Content format")
    parser.add_argument("--voice", default="formal", help="Voice style")
    parser.add_argument("--business_id", default="BAND_Foods", help="Business identifier")
    
    parser.add_argument("--verify-learning", action="store_true", help="Check learning system")
    parser.add_argument("--add-feedback", action="store_true", help="Add human feedback")
    parser.add_argument("--generation-id", help="Generation ID for feedback")
    parser.add_argument("--approved", type=str, choices=['yes', 'no'], help="Approved?")
    parser.add_argument("--score", type=float, help="Human score (0-10)")
    parser.add_argument("--feedback", help="Human feedback text")
    
    args = parser.parse_args()
    
    # Command: Verify learning
    if args.verify_learning:
        content_type = "blog" if "blog" in args.format.lower() else "social" if "social" in args.format.lower() else "ad"
        verify_learning_loop(args.business_id, content_type)
        exit(0)
    
    # Command: Add feedback
    if args.add_feedback:
        if not args.generation_id or not args.approved or args.score is None or not args.feedback:
            print("Error: --add-feedback requires --generation-id, --approved, --score, and --feedback")
            exit(1)
        
        human_approved = (args.approved.lower() == 'yes')
        update_generation_with_human_feedback(
            generation_id=args.generation_id,
            business_id=args.business_id,
            human_approved=human_approved,
            human_score=args.score,
            human_feedback=args.feedback
        )
        exit(0)
    
    # Default: Generate content
    print(f"\nStarting Iterative Content Generation")
    print(f"Topic: {args.topic}")
    print(f"Format: {args.format}")
    print(f"Business: {args.business_id}")
    print("="*60 + "\n")
    
    content_type = "blog" if "blog" in args.format.lower() else "social" if "social" in args.format.lower() else "ad"
    
    # Show learning summary
    learning_memory = BrandLearningMemory(business_id=args.business_id)
    print("\nLEARNING MEMORY SUMMARY")
    print("="*60)
    print(learning_memory.get_learning_summary(content_type))
    print("="*60 + "\n")
    
    # Generate with iterative refinement
    print("\n🔄 Starting iterative generation (up to 3 iterations)...\n")
    
    result, generation_id = run_iterative_generation(
        business_id=args.business_id,
        topic=args.topic,
        format_type=args.format,
        voice=args.voice
    )
    
    print("\n" + "="*60)
    print("GENERATION COMPLETE")
    print("="*60)
    print(f"\n📝 CONTENT:\n")
    print(result['content'])
    print("\n" + "="*60)
    print(f"Generation ID: {generation_id}")
    print(f"Final Score: {result['auto_score']:.1f}/10")
    print(f"Creative Angle: {result['creative_angle']}")
    print(f"Iterations Used: {result['iteration_count']}")
    
    if result.get('iteration_history'):
        print("\n📊 ITERATION HISTORY:")
        for iteration_info in result['iteration_history']:
            print(f"  Iteration {iteration_info['iteration']}: "
                  f"Score={iteration_info['score']:.1f}, "
                  f"Decision={iteration_info['decision']}, "
                  f"Issues={len(iteration_info.get('issues', []))}")
    
    print("="*60)
    
    # Offer to add feedback
    print("\n💬 To provide feedback on this generation, run:")
    print(f"python deployer.py --add-feedback \\")
    print(f"  --generation-id {generation_id} \\")
    print(f"  --business_id {args.business_id} \\")
    print(f"  --approved yes \\")
    print(f"  --score 9.0 \\")
    print(f"  --feedback 'Great work!'")
    print()